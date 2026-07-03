#!/usr/bin/env python3
"""Extract text snippets from priority project source files for impact assessment drafting."""

from __future__ import annotations

import argparse
import csv
from pathlib import Path

import openpyxl
import pdfplumber
from docx import Document


TEXT_SUFFIXES = {".txt", ".md", ".kml", ".ovkml", ".xml", ".csv"}
DOCX_SUFFIXES = {".docx"}
PDF_SUFFIXES = {".pdf"}
SHEET_SUFFIXES = {".xlsx", ".xlsm"}
PRIORITY_KEYWORDS = [
    "文物",
    "调查",
    "勘探",
    "验收",
    "意见",
    "函件",
    "备案",
    "红线",
    "坐标",
    "规划",
    "平面",
    "正射",
    "范围",
    "项目",
]
IGNORE_NAMES = {".DS_Store", "Thumbs.db"}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="提取项目资料优先文件的文本摘录")
    parser.add_argument("--项目资料目录", dest="project_dir", required=True)
    parser.add_argument("--输出文件", dest="output_file", required=True)
    parser.add_argument("--最多文件", dest="max_files", type=int, default=60)
    parser.add_argument("--每文件最大字符", dest="max_chars", type=int, default=2500)
    parser.add_argument("--pdf页数", dest="pdf_pages", type=int, default=5)
    return parser.parse_args()


def priority_score(path: Path) -> int:
    text = " ".join([path.name, *path.parts]).lower()
    score = 0
    for idx, keyword in enumerate(PRIORITY_KEYWORDS):
        if keyword.lower() in text:
            score += 100 - idx
    if path.suffix.lower() in DOCX_SUFFIXES | PDF_SUFFIXES | SHEET_SUFFIXES | TEXT_SUFFIXES:
        score += 30
    if "合同" in text or "商务" in text:
        score -= 60
    return score


def extract_docx(path: Path) -> str:
    doc = Document(path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables[:5]:
        for row in table.rows[:20]:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_pdf(path: Path, pages: int) -> str:
    texts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages[:pages]:
            text = page.extract_text() or ""
            if text.strip():
                texts.append(text.strip())
    if not texts:
        return "【未能直接提取文字，可能需要 OCR】"
    return "\n".join(texts)


def extract_sheet(path: Path) -> str:
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets[:3]:
        parts.append(f"## Sheet: {ws.title}")
        for row in ws.iter_rows(max_row=20, values_only=True):
            values = [str(value) for value in row if value is not None]
            if values:
                parts.append(" | ".join(values))
    return "\n".join(parts)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in DOCX_SUFFIXES:
        return extract_docx(path)
    if suffix in PDF_SUFFIXES:
        return extract_pdf(path, pages=ARGS.pdf_pages)
    if suffix in SHEET_SUFFIXES:
        return extract_sheet(path)
    if suffix == ".csv":
        rows = []
        with path.open("r", encoding="utf-8", errors="ignore") as f:
            for idx, row in enumerate(csv.reader(f)):
                if idx >= 30:
                    break
                rows.append(" | ".join(row))
        return "\n".join(rows)
    if suffix in TEXT_SUFFIXES:
        return path.read_text(encoding="utf-8", errors="ignore")
    return "【暂不支持该文件类型直接提取】"


def iter_priority_files(project_dir: Path, max_files: int) -> list[Path]:
    candidates = []
    allowed = DOCX_SUFFIXES | PDF_SUFFIXES | SHEET_SUFFIXES | TEXT_SUFFIXES
    for path in project_dir.rglob("*"):
        if not path.is_file() or path.name in IGNORE_NAMES:
            continue
        if path.suffix.lower() not in allowed:
            continue
        score = priority_score(path.relative_to(project_dir))
        if score > 0:
            candidates.append((score, path))
    candidates.sort(key=lambda item: (-item[0], str(item[1])))
    return [path for _, path in candidates[:max_files]]


def main() -> None:
    global ARGS
    ARGS = parse_args()
    project_dir = Path(ARGS.project_dir).expanduser().resolve()
    output_file = Path(ARGS.output_file).expanduser().resolve()
    if not project_dir.exists() or not project_dir.is_dir():
        raise SystemExit(f"项目资料目录不存在或不是文件夹：{project_dir}")
    files = iter_priority_files(project_dir, ARGS.max_files)
    output_file.parent.mkdir(parents=True, exist_ok=True)

    lines = [
        "# 项目资料文本摘录",
        "",
        f"项目资料目录：`{project_dir}`",
        f"提取文件数：{len(files)}",
        "",
    ]
    for path in files:
        rel = path.relative_to(project_dir)
        lines.extend([f"## {rel}", "", "```text"])
        try:
            text = extract_text(path).strip()
        except Exception as exc:  # noqa: BLE001
            text = f"【提取失败：{exc}】"
        if len(text) > ARGS.max_chars:
            text = text[: ARGS.max_chars] + "\n【后文截断】"
        lines.extend([text, "```", ""])
    output_file.write_text("\n".join(lines), encoding="utf-8")
    print(f"已提取 {len(files)} 个文件文本摘录：{output_file}")


if __name__ == "__main__":
    main()
