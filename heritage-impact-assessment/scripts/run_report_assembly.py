#!/usr/bin/env python3
"""Assemble v0.2.2 clean/evidence report drafts from pipeline outputs."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from pipeline_common import now_iso, read_json, read_jsonl, update_module_state, write_json, write_jsonl


FORMAL_HEADINGS = [
    "一、总则",
    "二、建设项目涉及文物概况",
    "三、建设项目规划概况",
    "四、项目用地范围与文物空间分布关系",
    "五、建设项目可能对文物造成的影响分析与评估",
    "六、减缓措施建议",
    "七、文物影响评估结论及建议",
    "八、支撑法律法规及文件",
]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="运行文物影响评估 v0.2.2 报告拼装阶段")
    parser.add_argument("--工作目录", "--workspace", dest="work_dir", required=True)
    parser.add_argument("--覆盖", dest="overwrite", action="store_true")
    return parser.parse_args()


def fact_value(rows: list[dict], field: str, default: str = "") -> str:
    for row in rows:
        if row.get("field_name") == field:
            return row.get("value") or default
    return default


FORBIDDEN_SUBMISSION_PATTERNS = [
    r"【待",
    r"待核对",
    r"待复核",
    r"待确认",
    r"待核验",
    r"待补充",
    r"未确认",
    r"未写明",
    r"资料缺失",
    r"后续补充",
    r"不确定",
    r"不编造坐标",
    r"正式报批阶段如主管部门要求",
    r"思考过程",
    r"推理过程",
    r"分析过程如下",
    r"OCR材料(?:显示|载明|表明)",
    r"保护规划修编OCR材料",
    r"调查材料(?:显示|载明|表明)",
    r"依据项目资料",
    r"项目资料(?:显示|载明|表明)",
    r"(?:可研|可行性研究报告)(?:从|称|提出|同时提出|显示|载明|表明)",
    r"根据《[^》]*(?:可研|可行性研究|建设方案|工程方案|勘探工作报告|文物调查|调查的报告|验收意见书|项目资料)[^》]*》",
    r"《[^》]*(?:可研|可行性研究|建设方案|工程方案|勘探工作报告|文物调查|调查的报告|验收意见书|项目资料)[^》]*》(中|内|显示|载明|指出|提出)",
]


def find_forbidden_submission_issues(markdown: str) -> list[str]:
    issues: list[str] = []
    for pattern in FORBIDDEN_SUBMISSION_PATTERNS:
        match = re.search(pattern, markdown)
        if match:
            excerpt_start = max(0, match.start() - 45)
            excerpt_end = min(len(markdown), match.end() + 45)
            excerpt = " ".join(markdown[excerpt_start:excerpt_end].split())
            issues.append(f"{pattern} -> {excerpt}")
    return issues


def sanitize_submission_text(text: str) -> str:
    replacements = [
        ("保护规划修编OCR材料显示，", ""),
        ("调查材料载明，", ""),
        ("可研同时提出，", "项目方案同时明确，"),
        ("可研提出", "项目方案提出"),
        ("可研称", "项目方案明确"),
        ("调查意见认为", "调查结论认为"),
    ]
    cleaned = text
    for old, new in replacements:
        cleaned = cleaned.replace(old, new)
    cleaned = re.sub(
        r"根据《[^》]*(?:可研|可行性研究|建设方案|工程方案|勘探工作报告|文物调查|调查的报告|验收意见书|项目资料)[^》]*》(?:和验收意见书)?，?",
        "",
        cleaned,
    )
    return cleaned.strip()


def md_table(rows: list[dict], columns: list[tuple[str, str]]) -> str:
    lines = ["| " + " | ".join(title for _, title in columns) + " |", "| " + " | ".join("---" for _ in columns) + " |"]
    for row in rows:
        values = [str(row.get(key, "")).replace("\n", " ") for key, _ in columns]
        lines.append("| " + " | ".join(values) + " |")
    return "\n".join(lines)


def regex_value(pattern: str, text: str) -> str:
    match = re.search(pattern, text)
    return match.group(1).strip() if match else ""


def skill_dir() -> Path:
    return Path(__file__).resolve().parents[1]


def markdown_tables_from_asset(relative_path: str) -> list[str]:
    path = skill_dir() / relative_path
    if not path.exists():
        return []
    tables: list[str] = []
    current: list[str] = []
    for line in path.read_text(encoding="utf-8").splitlines():
        if line.startswith("| "):
            current.append(line)
        elif current:
            tables.append("\n".join(current))
            current = []
    if current:
        tables.append("\n".join(current))
    return tables


def first_quote_for_section(quotes: list[dict], section_keyword: str, quote_type_keyword: str = "") -> dict:
    for quote in quotes:
        if section_keyword not in quote.get("target_section", ""):
            continue
        if quote_type_keyword and quote_type_keyword not in quote.get("quote_type", ""):
            continue
        return quote
    return {}


def project_basis_fixed_paragraph(project_name: str, heritage_name: str) -> str:
    return (
        f"现结合项目当前地理位置、施工方案，以及相关法律法规和行业标准，对{heritage_name}可能受到的影响进行评估。"
        f"具体来看，{project_name}在设计、建设、运营过程中，施工方式、车辆、人员等因素可能对不可移动文物本体及周边环境造成影响，"
        f"为统筹推进项目实施与文物保护，实现项目建设与文化遗产保护协同发展，需系统性评估{project_name}在建设及运营期间对所涉及文物的影响。"
    )


NECESSITY_FIXED_PARAGRAPHS = [
    "建设项目文物影响评估是对建设项目当前计划方案的实施对文物本体及其周边环境造成的影响进行分析、预测和评估，提出预防或减轻不良影响的对策和措施，并在规划和建设事中、事后进行跟踪监测。",
    "《中华人民共和国文物保护法》2024年11月8日修订版中第二十八条规定，在文物保护单位的保护范围内不得进行文物保护工程以外的其他建设工程或者爆破、钻探、挖掘等作业；因特殊情况需要进行的，必须保证文物保护单位的安全。因特殊情况需要在省级或者设区的市级、县级文物保护单位的保护范围内进行前款规定的建设工程或者作业的，必须经核定公布该文物保护单位的人民政府批准，在批准前应当征得上一级人民政府文物行政部门同意；在全国重点文物保护单位的保护范围内进行前款规定的建设工程或者作业的，必须经省、自治区、直辖市人民政府批准，在批准前应当征得国务院文物行政部门同意。",
    "第二十九条指出，在文物保护单位的建设控制地带内进行建设工程，不得破坏文物保护单位的历史风貌，工程设计方案应当根据文物保护单位的级别和建设工程对文物保护单位历史风貌的影响程度，经国家规定的文物行政部门同意后，依法取得建设工程规划许可。",
    "根据《内蒙古自治区人民政府办公厅关于加强工程建设文物保护前置审查工作的通知》（内政办发〔2024〕40号）、《内蒙古自治区文物局关于做好基本建设用地考古工作的通知》（内文物发〔2025〕6号）文件内容，对考古调查和勘探工作中发现文物遗存需原址保护的，或涉及各级文物保护单位保护范围和建设控制地带的建设用地，用地单位应提供《文物保护方案》和《文物影响评估报告》，由接收申请的文物行政部门组织包括自治区文物考古研究院专家在内的专家进行评审，并由参与评审的单位共同出具意见。",
]


CHARTER_CONVENTION_FIXED_LINES = [
    "(1) 《国际古迹保护与修复宪章》（1964）（International Charter for the Conservation and Restoration of Monuments and Sites）；",
    "(2) 《考古遗产保护与管理宪章》（1990）（Charter for the Protection and Management of the Archaeological Heritage）；",
    "(3) 《中国文物古迹保护准则》（2000）（Principles for the Conservation of Heritage Sites in China）；",
    "(4) 《会安草案——亚洲最佳保护范例》（2005）（Hoi An Protocols for Best Conservation of Historic Towns and Urban Areas）；",
    "(5) 《西安宣言》（2005）（Xi'an Declaration on the Conservation of the Setting of Heritage Structures, Sites and Areas）；",
    "(6) 《文化遗产阐释与展示宪章》（2008）（Charter on the Interpretation and Presentation of Cultural Heritage Sites）。",
]


LEGAL_REGULATION_FIXED_LINES = [
    "(1) 《中华人民共和国文物保护法》2024年11月8日修订，2025年3月1日起施行；",
    "(2) 《中华人民共和国文物保护法实施条例》2017年10月7日修订；",
    "(3) 《内蒙古自治区文物保护条例》2005年12月1日修订；",
]


POLICY_DOCUMENT_FIXED_LINES = [
    "(1)《国务院关于加强文化遗产保护的通知》（国发〔2005〕42号）；",
    "(2)《国务院关于进一步加强文物工作的指导意见》（国发〔2016〕17号）；",
    "(3)《关于加强基本建设工程中考古工作的指导意见》；",
    "(4)《中共中央办公厅、国务院办公厅关于加强文物保护利用改革的若干意见》（2018年10月8日）；",
    "(5)《内蒙古自治区人民政府关于进一步加强文物保护与利用工作的意见》（内政字〔2004〕260号）；",
    "(6)《内蒙古自治区人民政府办公厅关于加强工程建设文物保护前置审查工作的通知》（内政办发〔2024〕40号）；",
    "(7)《内蒙古自治区文物局关于做好基本建设用地考古工作的通知》（内文物发〔2025〕6号）。",
]


CHAPTER_EIGHT_FIXED_PARAGRAPHS = [
    "1.《中华人民共和国文物保护法》2024年11月8日修订，2025年3月1日起施行",
    "第二十八条 在文物保护单位的保护范围内不得进行文物保护工程以外的其他建设工程或者爆破、钻探、挖掘等作业；因特殊情况需要进行的，必须保证文物保护单位的安全。因特殊情况需要在省级或者设区的市级、县级文物保护单位的保护范围内进行前款规定的建设工程或者作业的，必须经核定公布该文物保护单位的人民政府批准，在批准前应当征得上一级人民政府文物行政部门同意；在全国重点文物保护单位的保护范围内进行前款规定的建设工程或者作业的，必须经省、自治区、直辖市人民政府批准，在批准前应当征得国务院文物行政部门同意。",
    "第二十九条 在文物保护单位的建设控制地带内进行建设工程，不得破坏文物保护单位的历史风貌，工程设计方案应当根据文物保护单位的级别和建设工程对文物保护单位历史风貌的影响程度，经国家规定的文物行政部门同意后，依法取得建设工程规划许可。",
    "2.《内蒙古自治区文物保护条例（修正）》",
    "第十四条 根据保护文物的实际需要，经自治区人民政府批准，可以在文物保护单位的周围划出一定的建设控制地带。在建设控制地带兴建建筑物，其设计方案，按文物保护单位的级别，在征得同级文化行政管理部门同意后，报城乡建设规划部门批准。",
    "第二十一条 在进行其他基本建设和生产建设时，任何单位或者个人发现文物，应负责保护好现场，并立即报告当地文化行政管理部门。当地文化行政管理部门应及时将情况报告上级直至自治区文化行政管理部门。",
    "3.《国务院关于加强文化遗产保护的通知》（国发〔2005〕42号）",
    "第三点第二条中明确指出：“严格执行重大建设工程项目审批、核准和备案制度。凡涉及文物保护事项的基本建设项目，必须依法在项目批准前征求文物行政部门的意见，在进行必要的考古勘探、发掘并落实文物保护措施以后方可实施。基本建设项目中的考古发掘要充分考虑文物保护工作的实际需要，加强统一管理，落实审批和监督责任。”",
    "4.《关于加强基本建设工程中考古工作的指导意见》（2007）",
    "在开展基本建设工程的“项目建议书”阶段，由文物考古机构收集建设项目涉及和影响区域内文物分布情况，提出初步文物保护意见，报省级文物行政部门确认后向设计单位提交《文物影响评估报告》。",
    "在工程建设的“可行性研究”阶段，由省级文物行政部门组织文物考古机构，对建设项目涉及和影响区域进行专项考古调查，编制《文物调查工作报告》，报省级文物行政部门认可后提交设计单位或建设单位。",
    "文物影响评估是由文物考古单位依据已掌握的资料，对建设项目涉及和影响区域内文物与建设工程的相互影响做出的分析评估。",
    "《文物影响评估报告》的内容应包括：建设项目涉及和影响区域内已有文物普查资料成果，已公布为各级文物保护单位保护范围和建设控制地带的相关资料，对项目选址及设计方案的初步建议。涉及省级以上文物保护单位的应报国家文物局。",
    "5.《内蒙古自治区人民政府办公厅关于加强工程建设文物保护前置审查工作的通知》（内政办发〔2024〕40号）",
    "第二条明确指出：全区各级文物行政部门应向社会公开各级文物保护单位保护范围和建设控制地带信息。实施大型基本建设工程，建设单位应事先报请自治区文物行政部门组织从事考古发掘的单位在工程选址范围内进行考古调查、勘探，由自治区文物行政部门在法定时限内出具审核意见。对城镇规划建设用地范围内以及各类开发区、高新区、产业集聚区、城乡一体化示范区等规划成片开发的土地开展区域评估时，应邀请自治区文物行政部门参与。经考古调查和勘探未发现文物的，自治区文物行政部门应在法定时限内出具审核意见；发现文物的，应当由自治区文物行政部门在勘探工作的基础上提出发掘计划，报国务院文物行政部门批准。确需在文物保护单位的保护范围和建设控制地带内进行建设工程的，按照《中华人民共和国文物保护法》有关规定执行。",
    "6.《内蒙古自治区文物局关于做好基本建设用地考古工作的通知》（内文物发〔2025〕6号）",
    "第二点第二条明确指出：对于考古调查和勘探工作中发现文物遗存需原址保护的，或涉及各级文物保护单位保护范围和建设控制地带的建设用地，用地单位应提供《文物保护方案》和《文物影响评估报告》，由接收申请的文物行政部门组织包括自治区文物考古研究院专家在内的专家进行评审，并由参与评审的单位共同出具意见。",
]


OTHER_MATERIAL_FIXED_LINES = [
    "(1) 《世界文化遗产影响评估指南》（2011年）；",
    "(2) 《中国文物古迹保护准则》（2015年）；",
]


OTHER_MATERIAL_INCLUDE_KEYWORDS = ["报告", "规划"]
OTHER_MATERIAL_EXCLUDE_KEYWORDS = [
    "图纸",
    "图片",
    "图件",
    "照片",
    "影像",
    "正射",
    "平面图",
    "红线",
    "KML",
    "OVKML",
    "kml",
    "ovkml",
    "坐标",
    "验收意见",
    "验收意见书",
]


def supplemental_legal_regulation_lines(project_rows: list[dict], heritage_rows: list[dict], quotes: list[dict]) -> list[str]:
    combined_text = "\n".join(
        str(value)
        for row in [*project_rows, *heritage_rows, *quotes]
        for value in row.values()
        if value
    )
    lines: list[str] = []
    if "长城" in combined_text or "长城保护条例" in combined_text:
        lines.append("(4) 《长城保护条例》2006年12月1日起施行。")
    return lines


def normalize_material_title(source_file: str) -> str:
    title = Path(source_file).stem.strip()
    return title.strip("《》")


def is_other_material_candidate(source_file: str) -> bool:
    title = normalize_material_title(source_file)
    if not title:
        return False
    if any(keyword in source_file or keyword in title for keyword in OTHER_MATERIAL_EXCLUDE_KEYWORDS):
        return False
    return any(keyword in title for keyword in OTHER_MATERIAL_INCLUDE_KEYWORDS)


def other_material_sort_key(title: str) -> tuple[int, str]:
    if "保护规划" in title:
        return (10, title)
    if "可行性研究报告" in title or "可研" in title:
        return (20, title)
    if "考古" in title and ("调查" in title or "勘探" in title) and "报告" in title:
        return (30, title)
    if "规划" in title:
        return (40, title)
    return (50, title)


def other_material_lines(text_index: list[dict]) -> list[str]:
    titles: list[str] = []
    seen: set[str] = set()
    for row in text_index:
        title = normalize_material_title(str(row.get("source_file", "")))
        if not is_other_material_candidate(str(row.get("source_file", ""))):
            continue
        if title in seen:
            continue
        seen.add(title)
        titles.append(title)

    titles.sort(key=other_material_sort_key)
    lines = [*OTHER_MATERIAL_FIXED_LINES]
    for offset, title in enumerate(titles, start=3):
        punctuation = "。" if offset == len(titles) + 2 else "；"
        lines.append(f"({offset}) 《{title}》{punctuation}")
    if len(lines) == len(OTHER_MATERIAL_FIXED_LINES):
        lines[-1] = lines[-1].rstrip("；") + "。"
    return lines


def evaluation_content_fixed_recommendation(project_name: str) -> str:
    return (
        f"根据文物影响的程度，明确{project_name}是否可行，提出调整建议和针对性的减缓措施。"
        f"并根据{project_name}的文物影响评估结论，提出其他相关建议。"
    )


def evaluation_method_fixed_blocks(project_name: str, heritage_name: str) -> list[tuple[str, str]]:
    return [
        ("heading", "1.调查法"),
        ("paragraph", "调查法是获取项目基本信息、现状资料、相关方意见及潜在影响的关键手段。本次评估将采用以下调查方式："),
        ("paragraph", "①资料收集调查"),
        ("paragraph", "系统性地搜集并整理与项目及文物相关的所有基础资料，具体包括下述内容："),
        ("paragraph", "文物信息。项目涉及文物的类型、等级、历史背景、保护范围、建设控制地带等详细资料。"),
        ("paragraph", "法律法规。梳理国家、自治区及地方层面涉及文物保护的法律法规、部门规章和行业标准，特别是《中华人民共和国文物保护法》《内蒙古自治区文物保护条例》及文物四有档案等核心文件。"),
        (
            "paragraph",
            f"项目信息。收集{project_name}的建设性质、具体内容、规模、范围、工程参数、相关审批文件、批复意见以及所在区域的城市或控制性规划。"
            "通过对上述资料的深入分析、统计和比较，初步判断该项目规划建设内容是否会对文物本体及其周围环境产生不利影响，为后续分析奠定基础，并为项目优化提供决策依据。",
        ),
        ("paragraph", "②现场调查"),
        (
            "paragraph",
            "访谈调查。主动走访并访谈相关的文物保护管理部门（如当地文物局或文物研究所）和项目建设单位。"
            "了解他们关于项目的计划安排、施工过程中可能运用的关键技术信息，以及对文物保护的意见和建议。"
            "同时，了解项目所在地及周边的自然环境、地质构造等基本情况。",
        ),
        (
            "paragraph",
            f"实地勘察。对{project_name}区域及{heritage_name}建设控制地带进行详细的实地踏勘。"
            "精确测量并记录建设项目范围与文物本体及保护范围的空间位置关系，核实项目当前的实际实施计划与申报计划的一致性，掌握第一手现场资料。"
            "确保评估工作基于真实、可靠的信息，访谈有助于理解各方的关切和专业技术要求，实地勘察则能直观、准确地掌握项目与文物的空间关系及现场实际状况，识别潜在风险点。",
        ),
        ("heading", "2.综合分析法"),
        ("paragraph", "在完成上述调查获取充分信息后，将运用以下综合分析方法对项目影响进行系统评估："),
        ("paragraph", "①清单对比法"),
        (
            "paragraph",
            "基于资料收集和现场调查的结果，分别整理项目规划建设要素和文物敏感要素及保护要求，如建设控制地带规定、环境要求等，"
            "将两类要素进行对比分析，以快速识别出项目与文物之间可能存在的直接冲突点和主要风险源。"
            "检查项目清单中的每一项要素是否与文物清单中的敏感要素或保护要求存在冲突或潜在干扰。"
            "例如，对比项目施工范围是否侵入文物保护范围或建设控制地带，对比项目采用的施工方法是否会产生对文物本体有害的振动或地下水扰动等。",
        ),
        ("paragraph", "②矩阵法"),
        (
            "paragraph",
            f"根据{project_name}的特点确立文物影响评价因子，识别出关键的影响路径和最敏感的文物要素，"
            "以系统化地评估项目各组成部分对文物各方面可能产生的综合影响，量化影响程度，识别出需要优先考虑和采取保护措施的关键环节。",
        ),
        ("paragraph", "③叠图法"),
        (
            "paragraph",
            "通过CAD制图、Sketch建模及GIS支持下的空间影像，分析项目与遗址文物保护规划的协调关系，与遗址遗迹的叠压关系，与环境景观和周边建筑的关系。"
            "识别出物理空间上的直接冲突或潜在威胁区域，为制定空间避让或工程防护措施提供依据。",
        ),
        ("heading", "3.预测法"),
        (
            "paragraph",
            f"依据项目建设工程的《{project_name}可行性研究报告》中的建设内容和施工方法，参考国家标准、行业标准，结合现场调查的各项数据，"
            "通过分析、类比、专业判断等方法预测建设内容对文物本体及周边历史环境的影响因素及程度、范围，为制定相关保护、减缓措施提供数据支持。",
        ),
    ]


def technical_spec_text(quotes: list[dict]) -> tuple[str, list[str]]:
    quote = first_quote_for_section(quotes, "4.技术规范", "技术规范原文摘录")
    if quote:
        evidence_ids = [quote.get("source_evidence_id", "")] if quote.get("source_evidence_id") else []
        return quote.get("text", ""), evidence_ids
    raise SystemExit("阻断：缺少可直接写入提交版的技术规范原文章节。请先补齐 quote_candidates.jsonl。")


def chapter_two_opening_lead(quotes: list[dict]) -> tuple[str, list[str]]:
    quote = first_quote_for_section(quotes, "二、建设项目涉及文物概况", "第二章开头总起段")
    if quote:
        evidence_ids = [quote.get("source_evidence_id", "")] if quote.get("source_evidence_id") else []
        return quote.get("text", ""), evidence_ids
    return "", []


def project_overview_natural_text(quotes: list[dict], project: list[dict]) -> tuple[str, list[str]]:
    quote = first_quote_for_section(quotes, "（一）项目概况", "项目概况自然段")
    if quote:
        evidence_ids = [quote.get("source_evidence_id", "")] if quote.get("source_evidence_id") else []
        return quote.get("text", ""), evidence_ids
    return (
        f"{fact_value(project, '项目名称')}位于{fact_value(project, '建设地点')}，建设单位为{fact_value(project, '建设单位')}。"
        f"项目主要建设内容为{fact_value(project, '建设内容')}，建设规模包括{fact_value(project, '项目面积')}。"
        "",
        ["E0001"],
    )


def project_necessity_blocks(quotes: list[dict]) -> tuple[list[str], list[str]]:
    quote = first_quote_for_section(quotes, "（二）项目建设必要性", "项目建设必要性原文摘录")
    if not quote:
        raise SystemExit("阻断：缺少项目建设必要性可研/建设方案原文章节，不能生成提交版。")
    evidence_ids = [quote.get("source_evidence_id", "")] if quote.get("source_evidence_id") else []
    text = quote.get("text", "").strip()
    blocks: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            blocks.append("")
            continue
        heading_match = re.match(r"^(\d+[.．、]\s*.+)$", line)
        if heading_match:
            blocks.append(f"#### {heading_match.group(1)}")
        else:
            blocks.append(line)
    return blocks or [text], evidence_ids


def quote_text_blocks(text: str) -> list[str]:
    blocks: list[str] = []
    for raw_line in text.strip().splitlines():
        line = raw_line.strip()
        if line:
            blocks.append(line)
        elif blocks and blocks[-1]:
            blocks.append("")
    while blocks and not blocks[-1]:
        blocks.pop()
    return blocks


def project_scheme_subsection_blocks(quotes: list[dict], section_keyword: str, quote_type_keyword: str, missing_text: str) -> tuple[list[str], list[str]]:
    quote = first_quote_for_section(quotes, section_keyword, quote_type_keyword)
    if not quote:
        return [missing_text], []
    evidence_ids = [quote.get("source_evidence_id", "")] if quote.get("source_evidence_id") else []
    return quote_text_blocks(quote.get("text", "")) or [quote.get("text", "")], evidence_ids


def site_selection_blocks(quotes: list[dict]) -> tuple[list[str], list[str]]:
    quote = first_quote_for_section(quotes, "选址分析", "选址分析原文摘录")
    if not quote:
        raise SystemExit("阻断：缺少选址分析/选址必要性分析可研或建设方案原文章节，不能生成提交版。")
    evidence_ids = [quote.get("source_evidence_id", "")] if quote.get("source_evidence_id") else []
    return quote_text_blocks(quote.get("text", "")) or [quote.get("text", "")], evidence_ids


def supporting_file_blocks(supporting_files: list[dict]) -> list[str]:
    if not supporting_files:
        raise SystemExit("阻断：缺少可插入第三章的立项、选址、核准文件或文物调查回函。")
    blocks: list[str] = []
    for index, row in enumerate(supporting_files, start=1):
        source_file = str(row.get("source_file", "")).strip()
        image_paths = row.get("image_paths") or []
        if image_paths:
            for image_index, image_path in enumerate(image_paths, start=1):
                blocks.append(f"![支持性文件正文第{image_index}页]({image_path})")
        elif source_file.lower().endswith(".pdf"):
            raise SystemExit(f"阻断：支持性文件 PDF 尚未转为正文页图片：{source_file}")
    return blocks


REGION_OVERVIEW_SUBSECTIONS = [
    ("1.地理位置", "【待联网检索：文物所在区域地理位置资料。优先使用政府网站、统计公报、地方志、文旅或文物部门公开资料，并标注来源。】"),
    ("2.自然概况", "【待联网检索：文物所在区域自然地理、地形地貌、气候水文、生态环境等资料。优先使用政府网站、统计公报、地方志、文旅或文物部门公开资料，并标注来源。】"),
    ("3.社会经济情况", "【待联网检索：文物所在区域人口、产业、交通、公共服务、社会经济发展等资料。优先使用政府网站、统计公报、地方志、文旅或文物部门公开资料，并标注来源。】"),
    ("4.历史沿革", "【待联网检索：文物所在区域历史沿革、历史文化背景、行政建置沿革和文化遗产背景等资料。优先使用政府网站、统计公报、地方志、文旅或文物部门公开资料，并标注来源。】"),
]


HERITAGE_OVERVIEW_SUBSECTIONS = [
    ("1.遗址基本情况", "【待依据文物调查回函或调查类材料补充：遗址基本情况，包括名称、级别、时代、类型、位置、历史背景和基本构成等。】"),
    ("2.遗址现状评估", "【待依据文物调查回函或调查类材料补充：遗址现状评估，包括保存现状、病害情况、周边环境和保护管理现状等。】"),
    ("3.调查、发掘、保护工程情况", "【待依据文物调查回函或调查类材料补充：调查、发掘、保护工程情况，包括历次调查、考古发掘、保护维修和相关工程情况等。】"),
    ("4.保护范围及建设控制地带", "【待依据文物调查回函或调查类材料补充：保护范围及建设控制地带，必须跟随材料原文术语和公布内容。】"),
    ("5.价值陈述", "【待依据文物调查回函或调查类材料补充：价值陈述，包括历史、科学、艺术、社会、景观等价值说明。】"),
]


VALUE_EVALUATION_FIXED_PARAGRAPH = (
    "《世界文化遗产影响评估指南》中的价值评估准则指出，对文物开展文物影响评估时，应考量其具备的文物价值、当地价值或国家价值，"
    "以及国家研究规划所明确的优先顺序与建议。同时，还需考虑文物所体现出的价值。"
    "本次文物影响评估报告依据《世界文化遗产影响评估指南》中的遗产价值分级量表（如下图所示），针对项目涉及的不可移动文物开展文物价值评估。"
)


def value_evaluation_replacement_paragraph(heritage: list[dict]) -> str:
    heritage_name = fact_value(heritage, "文物名称", "【文物名称】")
    heritage_level = fact_value(heritage, "文物级别", "【文物保护单位级别】")
    value_aspects = fact_value(heritage, "价值方面", "【价值方面】")
    heritage_value_grade = fact_value(heritage, "遗产价值分级", "【等级】")
    return (
        f"根据上述文物价值陈述，结合遗产价值分级量表，{heritage_name}属于已公布为保护单位的建筑物"
        f"（{heritage_level}），具备{value_aspects}。"
        f"因此评估{heritage_name}的遗产价值分级为“{heritage_value_grade}”。"
    )


def value_evaluation_fixed_blocks(heritage: list[dict]) -> list[str]:
    tables = markdown_tables_from_asset("assets/12-遗产价值分级量表模板.md")
    blocks = [VALUE_EVALUATION_FIXED_PARAGRAPH, "表1 遗产价值分级量表"]
    if tables:
        blocks.append(tables[0])
    else:
        blocks.append("【固定表格缺失：assets/12-遗产价值分级量表模板.md】")
    blocks.append(value_evaluation_replacement_paragraph(heritage))
    return blocks


def chapter_four_opening_lead(project_name: str, heritage_name: str) -> str:
    return (
        "本着“既有利于文物保护、又有利于基本建设”的“两利”方针，"
        f"在建设施工前对{project_name}涉及范围进行了现场勘察，"
        f"确定{heritage_name}与本项目的位置关系。"
    )


def chapter_four_first_section_intro(project: list[dict], heritage: list[dict], *, has_coordinates: bool) -> str:
    project_name = fact_value(project, "项目名称", "【项目名称】")
    location = fact_value(project, "建设地点", "【项目位置】")
    spatial_relation = fact_value(heritage, "空间关系", "【项目与文物空间关系】")
    suffix = "项目拐点坐标见表4。" if has_coordinates else ""
    return f"{project_name}位于{location}，{spatial_relation}。{suffix}"


def heritage_name_for_zoning_paragraph(heritage_name: str) -> str:
    if "桌子山秦长城东风农场" in heritage_name:
        return heritage_name.replace("桌子山秦长城东风农场", "桌子山秦长城——东风农场")
    return heritage_name


def chapter_four_zoning_fixed_paragraph(project_name: str, heritage_name: str) -> str:
    display_heritage = heritage_name_for_zoning_paragraph(heritage_name)
    return (
        "根据内蒙古自治区现有不可移动文物名录、内蒙古自治区长城资源调查数据库及现场调查情况，"
        f"{project_name}用地范围涉及乌海市海南区第五批自治区级文物保护单位{display_heritage}建设控制地带，不涉及其保护范围。"
    )


def chapter_five_heritage_name(heritage_name: str) -> str:
    return heritage_name.replace("桌子山秦长城——", "").replace("桌子山秦长城", "")


def chapter_five_opening_blocks(project_name: str, heritage_name: str) -> list[str]:
    display_heritage = chapter_five_heritage_name(heritage_name)
    return [
        "由于文化遗产的不可再生性，文物、遗产领域的法律法规及相关文件均反复指出，"
        "发展建设项目规划必须将文物保护纳入考量，通过履行规定的审批流程，确保项目建设尽可能规避对文物的破坏。"
        "建设单位应提供《文物保护方案》《文物影响评估报告》等文件，确保在施工前充分考虑到涉及的文物情况，"
        "避免建设活动导致相关文物与遗产的完整性的丧失，从而破坏遗产的真实性。",
        f"因此，需要对{project_name}可能对{display_heritage}建设控制地带产生的各类影响进行谨慎评估，"
        "并制定全面的保护措施，从而最大程度地减少项目在建设及运营过程中对不可移动文物本体及其周边环境的负面影响。"
        "本次评估涵盖项目方案阶段的合法合规性评估，以及施工及运营阶段的影响预测与评估。",
    ]


def chapter_five_integrated_impact_blocks(heritage: list[dict]) -> list[str]:
    heritage_name = fact_value(heritage, "文物名称", "【文物名称】")
    overall_result = fact_value(
        heritage,
        "文物影响评估总体结果",
        fact_value(heritage, "总体影响评估", "轻微"),
    )
    figure_no = fact_value(heritage, "影响评估量表图号", "9")
    table_no = fact_value(heritage, "综合影响表号", "9")
    image_path = skill_dir() / "assets/16-世界文化遗产影响评估指南-影响评估量表.png"
    return [
        "《世界文化遗产影响评估指南》提出，开发项目或其他改变对文化遗产属性的影响可能是正面的，也可能是负面的。"
        "因此有必要确认所有改变对遗产所有属性特征的影响，尤其是那些反映遗产突出普遍价值的属性。"
        "同时，还应确认某项改变对某一属性特征造成影响的规模和严重程度，这些组合在一起，就决定了遗产影响的重要性，或者称之为“效果的严重程度”。",
        f"本次文物影响评估依照《世界文化遗产影响评估指南》中的影响评估表（如图{figure_no}），"
        "以先前的文物价值评估结论为基础，综合考量项目对文物的影响内容及程度，进而开展文物影响评估工作。",
        f"![《世界文化遗产影响评估指南》影响评估量表]({image_path})",
        f"图{figure_no} 《世界文化遗产影响评估指南》影响评估量表",
        f"基于上述分析，结合{heritage_name}的遗产价值，依照《世界文化遗产影响评估指南》的影响评估量表，"
        f"对{heritage_name}的文物影响评估总体结果为“{overall_result}”（如表{table_no}所示）：",
    ]


def project_corner_coordinate_blocks(project_name: str, coordinates: list[dict]) -> list[str]:
    if not coordinates:
        return ["【缺少坐标表】"]
    return [
        f"表4 {project_name}拐点坐标表",
        md_table(coordinates, [("point_id", "序号"), ("longitude", "经度（E）"), ("latitude", "纬度（N）")]),
    ]


def mitigation_paragraph_blocks(mitigations: list[dict]) -> list[str]:
    if not mitigations:
        raise SystemExit("阻断：缺少第六章减缓措施矩阵，不能生成提交版。")
    grouped: dict[str, list[dict]] = {}
    for row in mitigations:
        phase = str(row.get("phase", "")).strip() or "全过程"
        grouped.setdefault(phase, []).append(row)
    phase_order = ["设计阶段", "施工期", "建设期", "运营期", "全过程"]
    ordered_phases = [phase for phase in phase_order if phase in grouped] + [
        phase for phase in grouped if phase not in phase_order
    ]
    blocks: list[str] = []
    chinese_numbers = "一二三四五六七八九十"
    for index, phase in enumerate(ordered_phases, start=1):
        number = chinese_numbers[index - 1] if index <= len(chinese_numbers) else str(index)
        blocks.append(f"### （{number}）{phase}减缓措施")
        for row in grouped[phase]:
            risk_source = str(row.get("risk_source", "")).strip()
            control_measure = str(row.get("control_measure", "")).strip()
            responsible_party = str(row.get("responsible_party", "")).strip()
            monitoring = str(row.get("monitoring_or_acceptance", "")).strip()
            sentence = control_measure
            if risk_source:
                sentence = f"针对{risk_source}，{sentence}"
            if responsible_party:
                sentence += f"相关措施由{responsible_party}落实。"
            if monitoring:
                sentence += f"实施效果通过{monitoring}进行检查。"
            blocks.append(sentence)
    return blocks


def exploration_source_text(work_dir: Path, text_index: list[dict]) -> str:
    chunks: list[str] = []
    for row in text_index:
        source_file = str(row.get("source_file", ""))
        if not ("勘探" in source_file or "考古调查" in source_file):
            continue
        text_path = row.get("text_path")
        if not text_path:
            continue
        path = work_dir / str(text_path)
        if path.exists():
            chunks.append(path.read_text(encoding="utf-8", errors="ignore"))
    return "\n".join(chunks)


def extract_exploration_company(text: str) -> str:
    organization = r"([^，。；;\n]{2,60}(?:公司|研究院|考古院|文物局|文物保护中心|博物馆|考古所|中心|院|所))"
    patterns = [
        rf"委托\s*{organization}\s*(?:开展|承担|实施|进行)",
        rf"由\s*{organization}\s*(?:开展|承担|实施|进行)",
    ]
    for pattern in patterns:
        value = regex_value(pattern, text)
        if value:
            return value
    return "【勘探公司】"


def extract_exploration_area(text: str, project: list[dict]) -> str:
    value = regex_value(r"(?:实际完成面积|勘探面积|考古调查、勘探面积)(?:为|[：:])?\s*([0-9.]+\s*(?:平方米|㎡))", text)
    if value:
        return value.replace("㎡", "平方米")
    return fact_value(project, "项目面积", "【勘探面积】")


def extract_exploration_conclusion(text: str) -> str:
    patterns = [
        r"在勘探区域内(未发现[^。；;\n]*(?:文物|文化层|遗迹)[^。；;\n]*)",
        r"在勘探区域内(发现[^。；;\n]*(?:文物|文化层|遗迹)[^。；;\n]*)",
    ]
    for pattern in patterns:
        value = regex_value(pattern, text)
        if value:
            return value.strip("。；;，, ")
    return "【勘探结论】"


def exploration_work_fixed_blocks(project_name: str, project: list[dict], text_index: list[dict], work_dir: Path) -> list[str]:
    text = exploration_source_text(work_dir, text_index)
    company = extract_exploration_company(text)
    area = extract_exploration_area(text, project)
    conclusion = extract_exploration_conclusion(text)
    return [
        "根据《内蒙古自治区人民政府办公厅关于加强工程建设文物保护前置审查工作的通知》（内政办发〔2024〕40号）、"
        "《内蒙古自治区文物局关于做好基本建设用地考古工作的通知》（内文物发〔2025〕6号）等文件规定，"
        "本项目开工建设前应聘请专业考古机构对拟建项目相关区域做进一步的考古调查、勘探工作。"
        f"因此，{project_name}委托{company}针对该项目区域展开了全面的考古调查勘探。"
        f"此次考古勘探聚焦于{project_name}用地范围。其核心目的在于为该区域开发提供地下文物遗存分布详细资料。"
        "借助系统的考古调查与勘探作业，详细查明建设区域内地下文物遗存的埋藏深度、分布范围、文化特质等关键基础信息，"
        "全力保障施工前地下文物的安全。与此同时，为项目建设施工环节准备科学严谨且翔实完备的考古资料，"
        "为该项目建设工程的后续顺利推进筑牢坚实基础，提供有力支撑与保障。",
        f"根据《{project_name}考古调查、勘探工作报告》，该次考古调查、勘探实际完成面积为{area}。"
        f"通过考古调查及勘探，认定在勘探区域内{conclusion}。",
    ]


def joined_row_text(rows: list[dict]) -> str:
    parts: list[str] = []
    for row in rows:
        for value in row.values():
            if value is None:
                continue
            parts.append(str(value))
    return " ".join(parts)


def chapter_seven_all_scheme_compliant(project: list[dict], design_compliance: list[dict], impacts: list[dict]) -> bool:
    for row in design_compliance:
        choice = str(row.get("chapter7_template_choice", "") or row.get("scheme_compliance_status", ""))
        if choice == "all_compliant":
            return True
        if choice == "compliant_after_adjustment":
            return False
        if choice == "blocking_unknown":
            raise SystemExit("第七章固定结论缺少已确认的项目方案合规状态，已阻断正文生成。")
    text = " ".join(
        [
            joined_row_text(design_compliance),
            fact_value(project, "项目方案合规性", ""),
            fact_value(project, "设计合规性", ""),
        ]
    )
    non_compliant_words = ["不符合", "不合规", "超过", "超出", "需调整", "须调整", "调整后", "不得高于", "不得超过"]
    if any(word in text for word in non_compliant_words):
        return False
    if design_compliance or "合规" in text or "符合" in text:
        return True
    impact_text = joined_row_text(impacts)
    return not any(word in impact_text for word in non_compliant_words)


def chapter_seven_zone_phrase(heritage: list[dict]) -> str:
    relation = fact_value(heritage, "空间关系", "")
    for zone in ["建设控制地带", "保护范围", "缓冲区", "五类缓冲区"]:
        if zone in relation:
            return zone
    zoning = fact_value(heritage, "保护范围及建设控制地带", "")
    for zone in ["建设控制地带", "保护范围", "缓冲区", "五类缓冲区"]:
        if zone in zoning:
            return zone
    return "建设控制地带"


def chapter_seven_control_requirement(requirements: list[dict], project: list[dict]) -> tuple[str, str]:
    for row in requirements:
        control_object = str(row.get("control_object", "") or row.get("控制对象", ""))
        control_requirement = str(row.get("control_requirement", "") or row.get("控制要求", ""))
        if control_object and control_requirement:
            return control_object, control_requirement
    text = " ".join(
        [
            joined_row_text(requirements),
            fact_value(project, "控制要求", ""),
            fact_value(project, "建筑高度控制要求", ""),
            fact_value(project, "建控地带限高", ""),
        ]
    )
    height = regex_value(r"不(?:得|应)?(?:超过|高于)\s*([0-9.]+\s*米)", text)
    if height:
        return "建筑高度", f"不超过{height}"
    density = regex_value(r"建筑密度不(?:得|应)?(?:超过|高于)\s*([0-9.]+\s*%)", text)
    if density:
        return "建筑密度", f"不超过{density}"
    if "风貌" in text:
        return "项目风貌", "与文物历史风貌相协调"
    return "相关建设内容", "文物保护相关要求范围内"


def chapter_seven_surrounding_text(project_name: str, heritage_name: str) -> str:
    if "寺" in heritage_name:
        return "寺院周边的环境"
    if any(keyword in project_name for keyword in ["厂", "加工", "园区", "基地"]):
        return "厂区周边的环境"
    return "项目周边环境"


def chapter_seven_activity_type(project_name: str) -> str:
    if any(keyword in project_name for keyword in ["寺", "宗教", "庙"]):
        return "宗教活动开展"
    if any(keyword in project_name for keyword in ["厂", "加工", "产业", "商业"]):
        return "经济效益增长"
    if any(keyword in project_name for keyword in ["展示", "旅游", "游客", "文化"]):
        return "文化展示利用"
    return "项目建设运营"


def chapter_seven_body_and_zoning_text(heritage: list[dict], zone_phrase: str) -> str:
    if "缓冲区" in zone_phrase:
        return "遗址本体及缓冲区"
    if "建设控制地带" in zone_phrase:
        return "文物本体及保护范围、建设控制地带"
    if "保护范围" in zone_phrase:
        return "文物本体及保护范围"
    return "文物本体及相关保护区域"


def chapter_seven_fixed_blocks(
    project_name: str,
    project: list[dict],
    heritage: list[dict],
    requirements: list[dict],
    design_compliance: list[dict],
    impacts: list[dict],
) -> list[tuple[str, str]]:
    heritage_name = fact_value(heritage, "文物名称", "相关文物")
    zone_phrase = chapter_seven_zone_phrase(heritage)
    control_object, control_requirement = chapter_seven_control_requirement([*design_compliance, *requirements], project)
    body_and_zoning = chapter_seven_body_and_zoning_text(heritage, zone_phrase)
    surrounding_text = chapter_seven_surrounding_text(project_name, heritage_name)
    activity_type = chapter_seven_activity_type(project_name)

    if chapter_seven_all_scheme_compliant(project, design_compliance, impacts):
        first_paragraph = (
            f"基于上述分析，本项目建设及运营过程对涉及的{heritage_name}的影响相对轻微。"
            "项目设计符合法律法规和相关规范，具有可行性。"
        )
    else:
        first_paragraph = (
            f"基于上述分析，本项目建设及运营过程对涉及的{heritage_name}{zone_phrase}的影响相对轻微。"
            f"但须全面落实保护措施，将{control_object}严格控制在{control_requirement}。"
            "调整后，项目设计符合法律法规和相关规范，具有可行性。"
        )

    return [
        ("heading", "（一）评估结论"),
        ("paragraph", first_paragraph),
        ("paragraph", "本次评估是依据现行的项目工程建设方案开展的，若建设方调整工程建设方案，需重新开展文物影响评估工作。"),
        (
            "paragraph",
            "通过系统地实施上述减缓措施，即强化环境监测、注重协调景观风貌、加强文物保护管理等多个维度，"
            f"能够将项目对{heritage_name}{zone_phrase}造成的不利影响控制在可接受限度内，"
            f"确保{body_and_zoning}的安全、稳定，以及历史环境的完整性。"
            f"建议项目单位构建长效的监测与评估机制，动态调整保护策略，以实现{activity_type}与文化遗产保护的平衡。",
        ),
        ("heading", "（二）其他建议"),
        (
            "paragraph",
            f"针对视觉景观、环境影响等问题，制定针对性的保护方案，确保{heritage_name}及其保护范围、建设控制地带的安全和完整。"
            f"建立监测体系，对{surrounding_text}进行定期监测，及时发现并解决问题。"
            "制定应急预案，应对可能出现的突发情况。与相关部门和单位加强沟通协调，共同做好文物保护工作。"
            "依据国家有关标准和规范，科学制定项目设计方案，确保文物保护单位及其环境的安全。"
            "项目施工前，应按照文物保护法律法规的规定和基本建设管理的相关要求，履行行政审批手续，"
            "经国家规定的文物行政部门同意后，依法取得建设工程规划许可。",
        ),
        (
            "paragraph",
            "若在施工过程中发现文物或遗存，应立即停工、保护现场，并报告文物行政主管部门，"
            "由文物行政主管部门依法组织处置，经批准后方可继续施工。",
        ),
    ]


def build_markdown(work_dir: Path, with_evidence: bool) -> tuple[str, list[dict]]:
    manifest = read_json(work_dir / "processing_output" / "manifest.json")
    project = read_jsonl(work_dir / "facts" / "project_facts.jsonl")
    heritage = read_jsonl(work_dir / "facts" / "heritage_facts.jsonl")
    requirements = read_jsonl(work_dir / "facts" / "requirement_facts.jsonl")
    quotes = read_jsonl(work_dir / "facts" / "quote_candidates.jsonl")
    impacts = read_jsonl(work_dir / "analysis" / "impact_matrix.jsonl")
    mitigations = read_jsonl(work_dir / "analysis" / "mitigation_matrix.jsonl")
    design_compliance = read_jsonl(work_dir / "analysis" / "chapter5_design_compliance.jsonl")
    figures = read_jsonl(work_dir / "processing_output" / "figure_index.jsonl")
    text_index = read_jsonl(work_dir / "processing_output" / "text_index.jsonl")
    supporting_files = read_jsonl(work_dir / "processing_output" / "supporting_files.jsonl")
    project_corner_coordinates = read_jsonl(work_dir / "processing_output" / "project_corner_coordinates.jsonl")

    project_name = manifest.get("project_name", "文物影响评估报告")
    evidence_suffix = "（证据版）" if with_evidence else ""
    lines = [f"# {project_name}文物影响评估报告{evidence_suffix}", ""]
    evidence_map = []

    def add(section: str, text: str, evidence_ids: list[str] | None = None) -> None:
        if not with_evidence:
            text = sanitize_submission_text(text)
        if not text:
            return
        lines.append(text + (f"【证据：{','.join(evidence_ids or [])}】" if with_evidence and evidence_ids else ""))
        lines.append("")
        evidence_map.append(
            {
                "section_id": section,
                "paragraph_id": f"P{len(evidence_map)+1:04d}",
                "text_excerpt": text[:120],
                "fact_ids": [],
                "evidence_ids": evidence_ids or [],
                "source_files": [],
                "confirmation_status": "待专业负责人确认" if "待确认" in text or "初步" in text else "待核验",
            }
        )

    lines.append(f"## {FORMAL_HEADINGS[0]}")
    lines.append("### （一）编制背景")
    lines.append("")
    lines.append("#### 1.评估项目基础信息")
    lines.append("")
    first_chapter_quote = first_quote_for_section(quotes, "一、总则", "评估项目基础信息")
    if first_chapter_quote:
        add(
            FORMAL_HEADINGS[0],
            first_chapter_quote.get("text", ""),
            [first_chapter_quote.get("source_evidence_id", "")] if first_chapter_quote.get("source_evidence_id") else [],
        )
    else:
        add(
            FORMAL_HEADINGS[0],
            f"{project_name}位于{fact_value(project, '建设地点')}，建设单位为{fact_value(project, '建设单位')}，"
            f"建设规模包括{fact_value(project, '项目面积')}、{fact_value(project, '总建筑面积')}、{fact_value(project, '总占地面积')}。",
            ["E0001"],
        )
    add(FORMAL_HEADINGS[0], project_basis_fixed_paragraph(project_name, fact_value(heritage, "文物名称")), ["E0001"])
    lines.append("#### 2.文物影响评估必要性")
    lines.append("")
    for paragraph in NECESSITY_FIXED_PARAGRAPHS:
        add(FORMAL_HEADINGS[0], paragraph)
    lines.append("### （三）评估内容")
    lines.append("")
    add(FORMAL_HEADINGS[0], evaluation_content_fixed_recommendation(project_name))
    lines.append("### （五）评估依据")
    lines.append("")
    lines.append("#### 1.宪章公约")
    lines.append("")
    for paragraph in CHARTER_CONVENTION_FIXED_LINES:
        add(FORMAL_HEADINGS[0], paragraph)
    lines.append("#### 2.法律法规")
    lines.append("")
    for paragraph in [*LEGAL_REGULATION_FIXED_LINES, *supplemental_legal_regulation_lines(project, heritage, quotes)]:
        add(FORMAL_HEADINGS[0], paragraph)
    lines.append("#### 3.文件规定")
    lines.append("")
    for paragraph in POLICY_DOCUMENT_FIXED_LINES:
        add(FORMAL_HEADINGS[0], paragraph)
    lines.append("#### 4.技术规范")
    lines.append("")
    text, evidence_ids = technical_spec_text(quotes)
    add(FORMAL_HEADINGS[0], text, evidence_ids)
    lines.append("#### 5.其他资料")
    lines.append("")
    for paragraph in other_material_lines(text_index):
        add(FORMAL_HEADINGS[0], paragraph)
    lines.append("### （六）评估方法")
    lines.append("")
    for block_type, text in evaluation_method_fixed_blocks(project_name, fact_value(heritage, "文物名称")):
        if block_type == "heading":
            lines.append(f"#### {text}")
            lines.append("")
        else:
            add(FORMAL_HEADINGS[0], text)

    lines.append(f"## {FORMAL_HEADINGS[1]}")
    lead_text, lead_evidence_ids = chapter_two_opening_lead(quotes)
    if lead_text:
        add(FORMAL_HEADINGS[1], lead_text, lead_evidence_ids)
    add(FORMAL_HEADINGS[1], f"项目涉及文物对象为{fact_value(heritage, '文物名称')}。{fact_value(heritage, '空间关系')}", ["E0001"])
    lines.append("### （一）文物所在区域概况")
    lines.append("")
    for heading, placeholder in REGION_OVERVIEW_SUBSECTIONS:
        lines.append(f"#### {heading}")
        lines.append("")
        add(FORMAL_HEADINGS[1], placeholder)
    lines.append(f"### （二）{fact_value(heritage, '文物名称')}概述")
    lines.append("")
    for heading, placeholder in HERITAGE_OVERVIEW_SUBSECTIONS:
        lines.append(f"#### {heading}")
        lines.append("")
        add(FORMAL_HEADINGS[1], placeholder)
    lines.append("#### 6.价值评估")
    lines.append("")
    for paragraph in value_evaluation_fixed_blocks(heritage):
        add(FORMAL_HEADINGS[1], paragraph)
    if quotes and not lead_text:
        fallback_text = quotes[-1].get("text", "")
        if not fallback_text:
            raise SystemExit("阻断：缺少可写入第二章的文物概况事实段。")
        add(FORMAL_HEADINGS[1], fallback_text, ["E0001"])

    lines.append(f"## {FORMAL_HEADINGS[2]}")
    lines.append("### （一）项目概况")
    lines.append("")
    overview_text, overview_evidence_ids = project_overview_natural_text(quotes, project)
    add(FORMAL_HEADINGS[2], overview_text, overview_evidence_ids)
    lines.append("### （二）项目建设必要性")
    lines.append("")
    necessity_blocks, necessity_evidence_ids = project_necessity_blocks(quotes)
    for block in necessity_blocks:
        if not block:
            lines.append("")
        elif block.startswith("#### "):
            lines.append(block)
            lines.append("")
        else:
            add(FORMAL_HEADINGS[2], block, necessity_evidence_ids)
    lines.append("### （三）支持性文件取得情况")
    lines.append("")
    for block in supporting_file_blocks(supporting_files):
        if block.startswith("!["):
            lines.append(block)
            lines.append("")
        else:
            add(FORMAL_HEADINGS[2], block)
    lines.append("### （四）建设及运营方案")
    lines.append("")
    for heading, section_keyword, quote_type_keyword, missing_text in [
        (
            "1.建设规模与技术指标",
            "1.建设规模与技术指标",
            "建设规模与技术指标原文摘录",
            "【待依据项目可研报告、可行性研究报告或建设方案完整摘取建设规模与技术指标相关内容；不得概述拟写。】",
        ),
        (
            "2.项目运营方案与规模",
            "2.项目运营方案与规模",
            "项目运营方案与规模原文摘录",
            "【待依据项目可研报告、可行性研究报告或建设方案完整摘取项目运营方案与规模相关内容；不得概述拟写。】",
        ),
    ]:
        lines.append(f"#### {heading}")
        lines.append("")
        scheme_blocks, scheme_evidence_ids = project_scheme_subsection_blocks(quotes, section_keyword, quote_type_keyword, missing_text)
        for block in scheme_blocks:
            if block:
                add(FORMAL_HEADINGS[2], block, scheme_evidence_ids)
            else:
                lines.append("")

    lines.append(f"## {FORMAL_HEADINGS[3]}")
    add(FORMAL_HEADINGS[3], chapter_four_opening_lead(project_name, fact_value(heritage, "文物名称")))
    lines.append("### （一）项目用地与文物相对位置关系")
    lines.append("")
    add(
        FORMAL_HEADINGS[3],
        chapter_four_first_section_intro(project, heritage, has_coordinates=bool(project_corner_coordinates)),
        ["E0001"],
    )
    for block in project_corner_coordinate_blocks(project_name, project_corner_coordinates):
        if block.startswith("| "):
            lines.append(block)
            lines.append("")
        else:
            add(FORMAL_HEADINGS[3], block)
    add(FORMAL_HEADINGS[3], chapter_four_zoning_fixed_paragraph(project_name, fact_value(heritage, "文物名称")))
    lines.append(f"### （二）{project_name}选址分析")
    lines.append("")
    selection_blocks, selection_evidence_ids = site_selection_blocks(quotes)
    for block in selection_blocks:
        if block:
            add(FORMAL_HEADINGS[3], block, selection_evidence_ids)
        else:
            lines.append("")
    lines.append("### （三）勘探工作")
    lines.append("")
    for block in exploration_work_fixed_blocks(project_name, project, text_index, work_dir):
        add(FORMAL_HEADINGS[3], block)

    lines.append(f"## {FORMAL_HEADINGS[4]}")
    for block in chapter_five_opening_blocks(project_name, fact_value(heritage, "文物名称")):
        add(FORMAL_HEADINGS[4], block)
    lines.append("### （一）项目方案设计合规性评估")
    lines.append("")
    chapter_five_tables = markdown_tables_from_asset("assets/13-第五章影响评估表模板.md")
    chapter_five_table_names = ["项目设计合规性评估", "项目建设期影响评估", "项目运营期影响评估"]
    for table_name, table in zip(chapter_five_table_names, chapter_five_tables):
        lines.append(f"### {table_name}")
        lines.append(table)
        lines.append("")
    lines.append("### （四）文物影响综合评估")
    lines.append("")
    for block in chapter_five_integrated_impact_blocks(heritage):
        if block.startswith("!["):
            lines.append(block)
            lines.append("")
        else:
            add(FORMAL_HEADINGS[4], block)
    annex_tables = markdown_tables_from_asset("assets/14-附表综合评估大表模板.md")
    if annex_tables:
        lines.append(annex_tables[0])
        lines.append("")
    else:
        add(FORMAL_HEADINGS[4], "【固定表格缺失：assets/14-附表综合评估大表模板.md】")

    lines.append(f"## {FORMAL_HEADINGS[5]}")
    for block in mitigation_paragraph_blocks(mitigations):
        if block.startswith("### "):
            lines.append(block)
            lines.append("")
        else:
            add(FORMAL_HEADINGS[5], block)

    lines.append(f"## {FORMAL_HEADINGS[6]}")
    for block_type, text in chapter_seven_fixed_blocks(project_name, project, heritage, requirements, design_compliance, impacts):
        if block_type == "heading":
            lines.append(f"### {text}")
            lines.append("")
        else:
            add(FORMAL_HEADINGS[6], text)

    lines.append(f"## {FORMAL_HEADINGS[7]}")
    for paragraph in CHAPTER_EIGHT_FIXED_PARAGRAPHS:
        add(FORMAL_HEADINGS[7], paragraph)

    annex_tables = markdown_tables_from_asset("assets/14-附表综合评估大表模板.md")
    if annex_tables:
        lines.append("## 附表")
        lines.append(f"### 附表一：{fact_value(heritage, '文物名称')}文物影响评估综合表")
        lines.append(annex_tables[0])
        lines.append("")

    return "\n".join(lines).strip() + "\n", evidence_map


def split_markdown_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator_row(cells: list[str]) -> bool:
    return all(cell.replace("-", "").replace(":", "").strip() == "" for cell in cells)


def set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 100, start: int = 100, bottom: int = 100, end: int = 100) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def format_table(table) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.shared import Pt

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, "D9EAF7")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = Pt(14)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(9)
                    if row_index == 0:
                        run.bold = True


def merge_integrated_assessment_header(table) -> None:
    if len(table.rows) < 2 or len(table.columns) < 10:
        return
    first_row = [cell.text.strip() for cell in table.rows[0].cells]
    if first_row[:7] != ["项目", "项目", "影响因子", "影响因子", "影响因子", "影响因子", "影响因子"]:
        return
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 0).text = "项目"
    table.cell(0, 2).merge(table.cell(0, 6))
    table.cell(0, 2).text = "影响因子"
    format_table(table)


def add_markdown_table(doc, table_lines: list[str]) -> None:
    rows = [split_markdown_table_row(line) for line in table_lines]
    rows = [row for row in rows if not is_separator_row(row)]
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=column_count)
    for row_index, row in enumerate(rows):
        for col_index in range(column_count):
            table.cell(row_index, col_index).text = row[col_index] if col_index < len(row) else ""
    format_table(table)
    merge_integrated_assessment_header(table)
    doc.add_paragraph()


def configure_doc_styles(doc) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(20)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for style_name, size, centered in [("Heading 1", 16, True), ("Heading 2", 14, False), ("Heading 3", 12, False)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        style.paragraph_format.line_spacing = Pt(20)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)
        if centered:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def write_docx(path: Path, markdown: str) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    doc = Document()
    configure_doc_styles(doc)
    lines = markdown.splitlines()
    title = lines[0][2:] if lines and lines[0].startswith("# ") else path.stem

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(180)
    title_run = cover.add_run(title)
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    title_run.font.size = Pt(22)
    doc.add_page_break()

    toc = doc.add_paragraph("目录")
    toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc.runs[0].bold = True
    toc.runs[0].font.size = Pt(16)
    doc.add_page_break()

    table_buffer: list[str] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            add_markdown_table(doc, table_buffer)
            table_buffer = []

    for line in lines[1:]:
        if line.startswith("| "):
            table_buffer.append(line)
            continue
        flush_table()
        image_match = re.match(r"!\[[^\]]*\]\(([^)]+)\)", line.strip())
        if image_match:
            image_path = Path(image_match.group(1))
            candidates = [image_path]
            if not image_path.is_absolute():
                candidates = [path.parent / image_path, skill_dir() / image_path]
            existing = next((candidate for candidate in candidates if candidate.exists()), None)
            if not existing:
                raise SystemExit(f"阻断：正文图片不存在，不能生成提交版：{image_match.group(1)}")
            picture = doc.add_paragraph()
            picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
            picture.add_run().add_picture(str(existing), width=Inches(5.8))
            continue
        if line.startswith("# "):
            continue
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif not line.strip():
            continue
        else:
            paragraph = doc.add_paragraph(line)
            paragraph.paragraph_format.first_line_indent = Pt(24)
    flush_table()
    doc.save(path)


def main() -> None:
    args = parse_args()
    work_dir = Path(args.work_dir).expanduser().resolve()
    clean_md, evidence_map = build_markdown(work_dir, with_evidence=False)
    evidence_md, evidence_map_e = build_markdown(work_dir, with_evidence=True)
    submission_issues = find_forbidden_submission_issues(clean_md)
    if submission_issues:
        issue_lines = "\n".join(f"- {issue}" for issue in submission_issues)
        raise SystemExit(f"阻断：清洁版仍含提交版禁用写法，请先修正上游事实/分析或拼装规则。\n{issue_lines}")
    (work_dir / "report_clean.md").write_text(clean_md, encoding="utf-8")
    (work_dir / "report_with_evidence.md").write_text(evidence_md, encoding="utf-8")
    write_jsonl(work_dir / "report_evidence_map.jsonl", evidence_map + evidence_map_e)
    write_docx(work_dir / "report_clean.docx", clean_md)
    write_docx(work_dir / "report_with_evidence.docx", evidence_md)
    update_module_state(work_dir, "report_assembly")
    write_json(
        work_dir / "run_state" / "report_assembly.module_done.json",
        {
            "module_name": "report_assembly",
            "status": "completed",
            "started_at": now_iso(),
            "finished_at": now_iso(),
            "input_files": ["facts/project_facts.jsonl", "facts/heritage_facts.jsonl", "analysis/impact_matrix.jsonl", "analysis/mitigation_matrix.jsonl"],
            "output_files": ["report_clean.md", "report_with_evidence.md", "report_clean.docx", "report_with_evidence.docx", "report_evidence_map.jsonl"],
            "blocking_gaps_count": 0,
            "issues_count": 0,
            "next_prompt": "",
            "notes": "报告拼装阶段完成。该脚本按 assets 与 docx 成稿规则输出完整章法、固定评估表、综合附表和基础 Word 版式；图件自动选图、插图位置和视觉 QA 仍需继续复核。",
        },
    )
    print(f"报告拼装完成：{work_dir}")


if __name__ == "__main__":
    main()
