#!/usr/bin/env python3
"""Render a Markdown impact assessment draft to a simple editable DOCX."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


BODY_FONT_SIZE = Pt(12)  # 小四
CAPTION_FONT_SIZE = Pt(10.5)  # 五号
HEADING_1_SIZE = Pt(16)  # 三号
HEADING_2_SIZE = Pt(14)  # 四号
BODY_LINE_SPACING = Pt(20)
FIRST_LINE_INDENT = Pt(24)
NUMBERED_LIST_RE = re.compile(r"^\d+[.．、]\s*")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="将文物影响评估 Markdown 正文转换为 docx")
    parser.add_argument("--输入", dest="input_md", required=True)
    parser.add_argument("--输出", dest="output_docx", required=True)
    parser.add_argument("--允许固定表格入口", dest="allow_table_placeholders", action="store_true")
    parser.add_argument("--编制单位", dest="compile_unit", default="【编制单位】")
    parser.add_argument("--编制时间", dest="compile_time", default="【编制时间】")
    return parser.parse_args()


def set_east_asia_font(style_or_run, font_name: str = "宋体") -> None:
    r_pr = style_or_run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        from docx.oxml import OxmlElement

        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def set_run_font(run, size=BODY_FONT_SIZE, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = size
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    set_east_asia_font(run, "宋体")


def set_paragraph_body_format(paragraph, line_spacing=BODY_LINE_SPACING) -> None:
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = line_spacing


def ensure_caption_style(doc: Document) -> None:
    if "图表题注" in [style.name for style in doc.styles]:
        return
    style = doc.styles.add_style("图表题注", WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = CAPTION_FONT_SIZE
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 0, 0)
    set_east_asia_font(style, "宋体")
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = BODY_FONT_SIZE
    normal.font.color.rgb = RGBColor(0, 0, 0)
    set_east_asia_font(normal, "宋体")
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = BODY_LINE_SPACING

    heading_specs = {
        "Heading 1": (HEADING_1_SIZE, WD_ALIGN_PARAGRAPH.CENTER, None),
        "Heading 2": (HEADING_2_SIZE, None, None),
        "Heading 3": (BODY_FONT_SIZE, None, FIRST_LINE_INDENT),
        "Heading 4": (BODY_FONT_SIZE, None, FIRST_LINE_INDENT),
    }
    for style_name, (size, alignment, first_indent) in heading_specs.items():
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = size
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        set_east_asia_font(style, "宋体")
        if alignment is not None:
            style.paragraph_format.alignment = alignment
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        style.paragraph_format.line_spacing = BODY_LINE_SPACING
        if first_indent is not None:
            style.paragraph_format.first_line_indent = first_indent

    ensure_caption_style(doc)


def split_table(lines: list[str]) -> list[list[str]]:
    table = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        table.append(cells)
    return table


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    rows = split_table([line for line in lines if not set(line.replace("|", "").strip()) <= {"-", ":"}])
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                set_paragraph_body_format(paragraph)
                for run in paragraph.runs:
                    set_run_font(run)


def is_caption(text: str) -> bool:
    return (
        (text.startswith("图") or text.startswith("表"))
        and len(text) >= 2
        and (text[1].isdigit() or text[1] in "一二三四五六七八九十")
    )


def is_image_line(text: str) -> bool:
    return text.startswith("![") or text.startswith("【图件占位")


def add_cover_and_toc(doc: Document, title: str, compile_unit: str, compile_time: str) -> None:
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_before = Pt(220)
    title_run = title_paragraph.add_run(title)
    set_run_font(title_run, size=Pt(22), bold=True)

    unit_paragraph = doc.add_paragraph()
    unit_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    unit_paragraph.paragraph_format.space_before = Pt(300)
    unit_run = unit_paragraph.add_run(compile_unit)
    set_run_font(unit_run, size=BODY_FONT_SIZE)

    time_paragraph = doc.add_paragraph()
    time_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    time_run = time_paragraph.add_run(compile_time)
    set_run_font(time_run, size=BODY_FONT_SIZE)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_run = toc_title.add_run("目录")
    set_run_font(toc_run, size=HEADING_1_SIZE, bold=True)

    toc_note = doc.add_paragraph()
    toc_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = toc_note.add_run("【目录预留页：请在 Word 中插入或更新自动目录】")
    set_run_font(note_run, size=BODY_FONT_SIZE)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def render(
    input_md: Path,
    output_docx: Path,
    allow_table_placeholders: bool = False,
    compile_unit: str = "【编制单位】",
    compile_time: str = "【编制时间】",
) -> None:
    doc = Document()
    configure_styles(doc)
    text = input_md.read_text(encoding="utf-8")
    if "【固定表格入口】" in text and not allow_table_placeholders:
        raise SystemExit(
            "正文中仍有【固定表格入口】。请先插入并填写遗产价值分级量表、第五章影响评估表和附表综合评估大表；"
            "如仅调试模板，可添加 --允许固定表格入口。"
        )
    lines = text.splitlines()
    title = next((line[2:].strip() for line in lines if line.startswith("# ")), input_md.stem)
    add_cover_and_toc(doc, title, compile_unit, compile_time)
    table_buffer: list[str] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            add_markdown_table(doc, table_buffer)
            table_buffer = []

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            table_buffer.append(stripped)
            continue
        flush_table()
        if not stripped:
            doc.add_paragraph("")
        elif stripped.startswith("# "):
            continue
        elif stripped.startswith("## "):
            heading_text = stripped[3:].strip()
            paragraph = doc.add_heading(heading_text, level=1)
            if heading_text.startswith("附件"):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=3)
        elif stripped.startswith("- "):
            paragraph = doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            set_paragraph_body_format(paragraph)
            for run in paragraph.runs:
                set_run_font(run)
        elif NUMBERED_LIST_RE.match(stripped):
            paragraph = doc.add_paragraph(stripped, style="List Number")
            set_paragraph_body_format(paragraph)
            for run in paragraph.runs:
                set_run_font(run)
        elif is_caption(stripped):
            paragraph = doc.add_paragraph(stripped, style="图表题注")
            for run in paragraph.runs:
                set_run_font(run, size=CAPTION_FONT_SIZE, bold=True)
        elif is_image_line(stripped):
            paragraph = doc.add_paragraph(stripped)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            for run in paragraph.runs:
                set_run_font(run)
        else:
            paragraph = doc.add_paragraph(stripped)
            set_paragraph_body_format(paragraph)
            for run in paragraph.runs:
                set_run_font(run)
    flush_table()
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)


def main() -> None:
    args = parse_args()
    input_md = Path(args.input_md).expanduser().resolve()
    output_docx = Path(args.output_docx).expanduser().resolve()
    if not input_md.exists():
        raise SystemExit(f"输入 Markdown 不存在：{input_md}")
    render(
        input_md,
        output_docx,
        allow_table_placeholders=args.allow_table_placeholders,
        compile_unit=args.compile_unit,
        compile_time=args.compile_time,
    )
    print(f"已生成 docx：{output_docx}")
    print("提醒：请在封面加盖电子章，并在 Word 中插入或更新目录。")


if __name__ == "__main__":
    main()
