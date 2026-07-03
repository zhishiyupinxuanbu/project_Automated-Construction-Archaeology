#!/usr/bin/env python3
"""Run v0.2.2 material processing and write machine-readable stage outputs."""

from __future__ import annotations

import argparse
import csv
import datetime as dt
import json
import re
import shutil
import subprocess
import tempfile
import urllib.error
import urllib.request
from pathlib import Path
from typing import Any


IGNORE_NAMES = {".DS_Store", "Thumbs.db"}
TEXT_SUFFIXES = {".txt", ".md", ".kml", ".ovkml", ".xml", ".csv"}
DOCX_SUFFIXES = {".docx"}
PDF_SUFFIXES = {".pdf"}
SHEET_SUFFIXES = {".xlsx", ".xlsm", ".csv"}
FIGURE_SUFFIXES = {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".webp", ".bmp"}
OCR_ROOT = Path("/Users/drevan01/Desktop/OCR")
TEXT_CHUNK_SIZE = 6000
TEXT_CHUNK_OVERLAP = 400
SUPPORTING_FILE_ALLOWED_KEYWORDS = [
    "立项",
    "选址",
    "核准",
    "批复",
    "备案",
    "文物调查回函",
    "文物调查函",
    "调查回函",
    "文物调查的报告",
]
SUPPORTING_FILE_EXCLUDED_KEYWORDS = [
    "红线",
    "正射",
    "影像",
    "规划平面图",
    "总平面",
    "KML",
    "OVKML",
    "kml",
    "ovkml",
    "坐标",
    "考古调查",
    "勘探",
    "验收",
    "意见书",
    "文物影响评估",
    "文评",
    "保护方案",
    "评审",
    "图纸",
    "图片",
    "照片",
]

CATEGORY_RULES = [
    ("project_basic", "项目基本信息", ["项目", "建设", "改扩建", "可研", "建议书", "方案", "规划"], "三、建设项目规划概况"),
    ("heritage_basic", "文物基本信息", ["文物", "长城", "遗址", "保护单位", "调查"], "二、建设项目涉及文物概况"),
    ("protection_zoning", "保护区划信息", ["保护范围", "建设控制地带", "建控", "缓冲区", "管控"], "二、建设项目涉及文物概况；四、项目用地范围与文物空间分布关系"),
    ("spatial_relation", "空间关系信息", ["坐标", "kml", "ovkml", "红线", "范围", "正射", "位置", "距离", "宗地图"], "四、项目用地范围与文物空间分布关系"),
    ("design", "工程设计信息", ["总平", "平面图", "设计", "建筑", "施工图", "初设"], "三、建设项目规划概况；五、建设项目可能对文物造成的影响分析与评估"),
    ("survey", "调查勘探信息", ["考古", "勘探", "调查", "验收", "探孔", "剖线"], "四、项目用地范围与文物空间分布关系；五、建设项目可能对文物造成的影响分析与评估"),
    ("review", "审查评审信息", ["函件", "批复", "意见", "请示", "复函", "核查"], "一、总则；七、文物影响评估结论及建议"),
    ("figure", "图件照片信息", ["图", "影像", "照片", "示意", "正射", "位置"], "二、建设项目涉及文物概况；四、项目用地范围与文物空间分布关系；正文图件"),
    ("business", "商务资料", ["合同", "商务", "询价", "营业执照", "身份证"], "只登记不深处理"),
]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="运行文物影响评估 v0.2.2 资料处理阶段")
    parser.add_argument("--项目资料目录", dest="project_dir", required=True)
    parser.add_argument("--输出目录", dest="output_dir", required=True)
    parser.add_argument("--项目名称", dest="project_name", required=True)
    parser.add_argument("--覆盖", dest="overwrite", action="store_true")
    parser.add_argument("--每文件最大字符", dest="max_chars", type=int, default=0, help="0 表示不截断，默认全量抽取")
    return parser.parse_args()


def now_iso() -> str:
    return dt.datetime.now(dt.timezone.utc).astimezone().isoformat(timespec="seconds")


def write_json(path: Path, data: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def append_jsonl(path: Path, row: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as f:
        f.write(json.dumps(row, ensure_ascii=False) + "\n")


def limit_text(text: str, max_chars: int) -> str:
    if max_chars and max_chars > 0:
        return text[:max_chars]
    return text


def write_text_chunks(text: str, *, file_id: str, source_file: str, text_dir: Path, output_dir: Path, index_path: Path) -> int:
    if not text:
        return 0
    chunk_dir = text_dir / f"{file_id}_chunks"
    chunk_dir.mkdir(parents=True, exist_ok=True)
    start = 0
    chunk_count = 0
    text_len = len(text)
    step = max(1, TEXT_CHUNK_SIZE - TEXT_CHUNK_OVERLAP)
    while start < text_len:
        end = min(text_len, start + TEXT_CHUNK_SIZE)
        chunk = text[start:end]
        chunk_count += 1
        chunk_path = chunk_dir / f"{file_id}_chunk_{chunk_count:04d}.txt"
        chunk_path.write_text(chunk, encoding="utf-8")
        append_jsonl(
            index_path,
            {
                "chunk_id": f"{file_id}-CH{chunk_count:04d}",
                "file_id": file_id,
                "source_file": source_file,
                "chunk_path": str(chunk_path.relative_to(output_dir)),
                "char_start": start,
                "char_end": end,
                "char_count": len(chunk),
            },
        )
        if end >= text_len:
            break
        start += step
    return chunk_count


def classify(rel: Path) -> list[tuple[str, str, str]]:
    haystack = " ".join([rel.name, *rel.parts]).lower()
    matches: list[tuple[str, str, str]] = []
    for code, label, keywords, section in CATEGORY_RULES:
        if any(keyword.lower() in haystack for keyword in keywords):
            matches.append((code, label, section))
    suffix = rel.suffix.lower()
    if suffix in FIGURE_SUFFIXES and not any(code == "figure" for code, _, _ in matches):
        matches.append(("figure", "图件照片信息", "正文图件"))
    if suffix in {".kml", ".ovkml"} and not any(code == "spatial_relation" for code, _, _ in matches):
        matches.append(("spatial_relation", "空间关系信息", "四、项目用地范围与文物空间分布关系"))
    if not matches:
        matches.append(("other", "其他", "待判断"))
    return matches


def is_supporting_file_candidate(rel: Path) -> bool:
    haystack = " ".join([rel.name, *rel.parts])
    if any(keyword in haystack for keyword in SUPPORTING_FILE_EXCLUDED_KEYWORDS):
        return False
    return any(keyword in haystack for keyword in SUPPORTING_FILE_ALLOWED_KEYWORDS)


def supporting_file_type(rel: Path) -> str:
    haystack = " ".join([rel.name, *rel.parts])
    if any(keyword in haystack for keyword in ["文物调查回函", "文物调查函", "调查回函", "文物调查的报告"]):
        return "文物调查回函"
    return "立项、选址或核准文件"


def safe_stem(text: str) -> str:
    safe = re.sub(r"[^\w\u4e00-\u9fff-]+", "_", text).strip("_")
    return safe or "supporting_file"


def convert_supporting_pdf_to_images(path: Path, *, file_id: str, rel: Path, image_dir: Path, output_dir: Path) -> tuple[list[str], str | None]:
    if not shutil.which("pdftoppm"):
        return [], "pdftoppm 不可用，支持性 PDF 未能转为正文图片"
    target_dir = image_dir / file_id
    target_dir.mkdir(parents=True, exist_ok=True)
    prefix = target_dir / safe_stem(Path(rel).stem)
    try:
        result = subprocess.run(
            ["pdftoppm", "-png", "-r", "180", str(path), str(prefix)],
            text=True,
            capture_output=True,
            timeout=120,
        )
    except Exception as exc:
        return [], f"支持性 PDF 转图片失败：{exc}"
    if result.returncode != 0:
        return [], f"支持性 PDF 转图片失败：{result.stderr.strip()}"
    image_paths = sorted(target_dir.glob(f"{prefix.name}-*.png"))
    if not image_paths:
        return [], "支持性 PDF 转图片未生成图片"
    return [str(item.relative_to(output_dir)) for item in image_paths], None


def is_project_corner_coordinate_source(rel: Path) -> bool:
    haystack = " ".join([rel.name, *rel.parts])
    if not any(keyword in haystack for keyword in ["拐点", "经纬度", "坐标"]):
        return False
    return not any(keyword in haystack for keyword in ["文物", "文保", "长城", "保护范围", "建设控制地带"])


def normalize_header(value: Any) -> str:
    return str(value or "").strip().replace(" ", "").replace("\n", "")


def coordinate_rows_from_records(records: list[list[Any]], *, file_id: str, source_file: str) -> list[dict[str, Any]]:
    header_index = -1
    header_map: dict[str, int] = {}
    for index, row in enumerate(records):
        headers = [normalize_header(value) for value in row]
        point_col = next((i for i, value in enumerate(headers) if value in {"序号", "点号", "拐点", "拐点编号", "编号"}), -1)
        longitude_col = next((i for i, value in enumerate(headers) if "经度" in value or value.upper() in {"E", "LON", "LONGITUDE"}), -1)
        latitude_col = next((i for i, value in enumerate(headers) if "纬度" in value or value.upper() in {"N", "LAT", "LATITUDE"}), -1)
        if point_col >= 0 and longitude_col >= 0 and latitude_col >= 0:
            header_index = index
            header_map = {"point": point_col, "longitude": longitude_col, "latitude": latitude_col}
            break
    if header_index < 0:
        return []
    rows: list[dict[str, Any]] = []
    for row_number, row in enumerate(records[header_index + 1 :], start=header_index + 2):
        values = [str(value or "").strip() for value in row]
        max_index = max(header_map.values())
        if len(values) <= max_index:
            continue
        point_id = values[header_map["point"]]
        longitude = values[header_map["longitude"]]
        latitude = values[header_map["latitude"]]
        if not point_id or not longitude or not latitude:
            continue
        rows.append(
            {
                "coordinate_id": f"{file_id}-C{len(rows)+1:03d}",
                "file_id": file_id,
                "source_file": source_file,
                "row_number": row_number,
                "point_id": point_id,
                "longitude": longitude,
                "latitude": latitude,
            }
        )
    return rows


def extract_project_corner_coordinates(path: Path, *, file_id: str, rel: Path) -> tuple[list[dict[str, Any]], str | None]:
    source_file = str(rel)
    suffix = path.suffix.lower()
    try:
        if suffix in {".xlsx", ".xlsm"}:
            from openpyxl import load_workbook

            workbook = load_workbook(path, data_only=True, read_only=True)
            records: list[list[Any]] = []
            for sheet in workbook.worksheets:
                records.extend([list(row) for row in sheet.iter_rows(values_only=True)])
            return coordinate_rows_from_records(records, file_id=file_id, source_file=source_file), None
        if suffix == ".csv":
            with path.open("r", encoding="utf-8-sig", errors="ignore", newline="") as f:
                return coordinate_rows_from_records(list(csv.reader(f)), file_id=file_id, source_file=source_file), None
    except Exception as exc:
        return [], f"项目拐点坐标表解析失败：{exc}"
    return [], None


def collect_files(project_dir: Path) -> list[Path]:
    files: list[Path] = []
    for path in sorted(project_dir.rglob("*")):
        if not path.is_file() or path.name in IGNORE_NAMES:
            continue
        files.append(path)
    return files


def read_text_file(path: Path, max_chars: int) -> str:
    return limit_text(path.read_text(encoding="utf-8", errors="ignore"), max_chars)


def extract_docx(path: Path, max_chars: int) -> tuple[str, str | None, str]:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - depends on environment
        return "", f"python-docx 不可用：{exc}", "docx"
    try:
        doc = Document(path)
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables[:5]:
            for row in table.rows[:20]:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return limit_text("\n".join(parts), max_chars), None, "docx"
    except Exception as exc:
        return "", f"DOCX 提取失败：{exc}", "docx"


def extract_pdf_with_pdfplumber(path: Path, max_chars: int) -> tuple[str, str | None]:
    try:
        import pdfplumber
    except Exception as exc:  # pragma: no cover - depends on environment
        return "", f"pdfplumber 不可用，可能需要 OCR：{exc}"
    try:
        texts: list[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                if text.strip():
                    texts.append(text.strip())
                if max_chars and max_chars > 0 and sum(len(item) for item in texts) >= max_chars:
                    break
        if not texts:
            return "", "PDF 未能直接提取文字，可能需要 OCR"
        return limit_text("\n".join(texts), max_chars), None
    except Exception as exc:
        return "", f"PDF 提取失败：{exc}"


def extract_pdf_with_pdftotext(path: Path, max_chars: int) -> tuple[str, str | None]:
    if not shutil.which("pdftotext"):
        return "", "pdftotext 不可用"
    try:
        result = subprocess.run(
            ["pdftotext", "-layout", str(path), "-"],
            text=True,
            capture_output=True,
            timeout=120,
        )
    except Exception as exc:
        return "", f"pdftotext 提取失败：{exc}"
    if result.returncode != 0:
        return "", f"pdftotext 提取失败：{result.stderr.strip()}"
    text = result.stdout.strip()
    return (limit_text(text, max_chars), None) if text else ("", "pdftotext 未提取到文字")


def extract_pdf_with_pypdf(path: Path, max_chars: int) -> tuple[str, str | None]:
    try:
        from pypdf import PdfReader
    except Exception as exc:
        return "", f"pypdf 不可用：{exc}"
    try:
        reader = PdfReader(str(path))
        texts = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                texts.append(text.strip())
            if max_chars and max_chars > 0 and sum(len(item) for item in texts) >= max_chars:
                break
        if not texts:
            return "", "pypdf 未提取到文字"
        return limit_text("\n".join(texts), max_chars), None
    except Exception as exc:
        return "", f"pypdf 提取失败：{exc}"


def local_ocr_service_available() -> bool:
    try:
        with urllib.request.urlopen("http://127.0.0.1:8000/health", timeout=2) as response:
            return response.status == 200
    except (OSError, urllib.error.URLError):
        return False


def extract_pdf_with_local_ocr(path: Path, max_chars: int) -> tuple[str, str | None]:
    client = OCR_ROOT / "scripts" / "ocr_client.py"
    python = OCR_ROOT / ".venv" / "bin" / "python"
    if not client.exists() or not python.exists():
        return "", f"本地 Paddle OCR 工具不存在：{OCR_ROOT}"
    if not local_ocr_service_available():
        return "", "本地 Paddle OCR 服务未运行，需先启动 /Users/drevan01/Desktop/OCR/scripts/start_service.sh"
    try:
        with tempfile.TemporaryDirectory() as tmp:
            output_json = Path(tmp) / f"{path.stem}.ocr.json"
            result = subprocess.run(
                [str(python), str(client), str(path), "--output", str(output_json)],
                text=True,
                capture_output=True,
                timeout=600,
                cwd=str(OCR_ROOT),
            )
            if result.returncode != 0:
                return "", f"本地 Paddle OCR 调用失败：{result.stderr.strip()}"
            data = json.loads(output_json.read_text(encoding="utf-8"))
            text = str(data.get("text", "")).strip()
            return (limit_text(text, max_chars), None) if text else ("", "本地 Paddle OCR 未识别到文字")
    except Exception as exc:
        return "", f"本地 Paddle OCR 调用失败：{exc}"


def extract_pdf(path: Path, max_chars: int) -> tuple[str, str | None, str]:
    attempts: list[str] = []
    for method, extractor in [
        ("pdfplumber", extract_pdf_with_pdfplumber),
        ("pdftotext", extract_pdf_with_pdftotext),
        ("pypdf", extract_pdf_with_pypdf),
        ("paddle_ocr", extract_pdf_with_local_ocr),
    ]:
        text, issue = extractor(path, max_chars)
        if text:
            return text, None, method
        if issue:
            attempts.append(issue)
    return "", "；".join(attempts), "unresolved_pdf"


def extract_text(path: Path, max_chars: int) -> tuple[str, str | None, str]:
    suffix = path.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return read_text_file(path, max_chars), None, "direct"
    if suffix in DOCX_SUFFIXES:
        return extract_docx(path, max_chars)
    if suffix in PDF_SUFFIXES:
        return extract_pdf(path, max_chars)
    return "", None, "none"


def reset_output(output_dir: Path, overwrite: bool) -> None:
    if output_dir.exists() and any(output_dir.iterdir()):
        if not overwrite:
            raise SystemExit(f"输出目录非空，如需覆盖请添加 --覆盖：{output_dir}")
        shutil.rmtree(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)


def write_files_csv(path: Path, rows: list[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fields = [
        "file_id",
        "relative_path",
        "suffix",
        "size_bytes",
        "categories",
        "deep_process",
        "skip_reason",
        "applicable_sections",
    ]
    with path.open("w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fields})


def write_evidence_csv(path: Path, rows: list[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fields = [
        "evidence_id",
        "source_file_id",
        "source_file",
        "source_location",
        "evidence_type",
        "original_excerpt",
        "structured_value",
        "use_mode",
        "applicable_section",
        "confidence",
        "confirmation_reason",
    ]
    with path.open("w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fields})


def blocking_gaps(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    categories = set()
    for row in rows:
        for item in row["category_codes"]:
            categories.add(item)
    required = [
        ("project_basic", "项目基本事实", "缺少任何可替代项目基本信息来源"),
        ("heritage_basic", "文物对象身份", "缺少任何可替代文物对象来源"),
        ("spatial_relation", "空间边界", "缺少任何可替代空间关系、红线或坐标来源"),
    ]
    gaps = []
    for code, gap_type, reason in required:
        if code not in categories:
            gaps.append(
                {
                    "gap_id": f"G{len(gaps)+1:03d}",
                    "gap_type": gap_type,
                    "blocking_reason": reason,
                    "why_not_replaceable": "项目资料包中未识别到可替代来源；需人工确认是否可由长期知识库或官方联网来源补足。",
                    "required_user_action": f"补充或指定{gap_type}相关资料。",
                    "affected_downstream_step": "事实抽取、分析判断、报告拼装",
                }
            )
    return gaps


def write_next_prompt(path: Path, project_name: str, output_dir: Path, blocked: bool) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if blocked:
        path.write_text(
            "\n".join(
                [
                    "# 阻断：需要用户补充不可替代资料",
                    "",
                    f"项目名称：{project_name}",
                    f"项目工作包：{output_dir}",
                    "",
                    "请读取 `processing_output/user_blocking_gaps.jsonl`，只向用户列出不可替代资料缺项。",
                    "不要进入事实抽取、分析判断或报告拼装阶段。",
                ]
            )
            + "\n",
            encoding="utf-8",
        )
        return
    path.write_text(
        "\n".join(
            [
                "# 下一阶段启动提示",
                "",
                "## 阶段",
                "fact_extraction",
                "",
                "## 项目路径",
                f"`{output_dir}`",
                "",
                "## 必读规则文件",
                "- `references/03-事实抽取模块.md`",
                "- `references/01-证据台账模块.md`",
                "- `references/07-项目资料识别规则.md`",
                "- `references/08-事实抽取字段表.md`",
                "- `references/09-章节证据映射规则.md`",
                "- `references/11-固定正文结构与固定内容.md`",
                "",
                "## 只读输入",
                "- `module_state.json`",
                "- `run_state/material_processing.module_done.json`",
                "- `processing_output/manifest.json`",
                "- `processing_output/files.csv`",
                "- `processing_output/text_index.jsonl`",
                "- `processing_output/text_chunk_index.jsonl`",
                "- `processing_output/table_index.jsonl`",
                "- `processing_output/figure_index.jsonl`",
                "- `processing_output/supporting_files.jsonl`",
                "- `processing_output/external_sources.jsonl`",
                "- `evidence/evidence_register.jsonl`",
                "",
                "## 必写输出",
                "- `facts/project_facts.jsonl`",
                "- `facts/heritage_facts.jsonl`",
                "- `facts/quote_candidates.jsonl`",
                "- `facts/fact_issues.jsonl`",
                "- `run_state/fact_extraction.module_done.json`",
                "- `next_prompts/next_prompt_analysis.md`",
                "",
                "## 禁止事项",
                "- 不读取上一阶段聊天记录。",
                "- 不通读全部原始资料；按索引读取必要片段。",
                "- 不继承既有文评、园区评估或保护方案的结论。",
            ]
        )
        + "\n",
        encoding="utf-8",
    )


def main() -> None:
    args = parse_args()
    project_dir = Path(args.project_dir).expanduser().resolve()
    output_dir = Path(args.output_dir).expanduser().resolve()
    if not project_dir.is_dir():
        raise SystemExit(f"项目资料目录不存在或不是文件夹：{project_dir}")
    reset_output(output_dir, args.overwrite)

    started_at = now_iso()
    processing_dir = output_dir / "processing_output"
    text_dir = processing_dir / "extracted_text"
    text_chunk_dir = processing_dir / "extracted_text_chunks"
    table_dir = processing_dir / "extracted_tables"
    supporting_image_dir = processing_dir / "supporting_file_images"
    figure_dir = output_dir / "extracted_figures"
    evidence_dir = output_dir / "evidence"
    run_state_dir = output_dir / "run_state"
    prompt_dir = output_dir / "next_prompts"
    debug_dir = processing_dir / "debug"
    for directory in [text_dir, text_chunk_dir, table_dir, supporting_image_dir, figure_dir, evidence_dir, run_state_dir, prompt_dir, debug_dir]:
        directory.mkdir(parents=True, exist_ok=True)

    file_rows: list[dict[str, Any]] = []
    evidence_rows: list[dict[str, Any]] = []
    files = collect_files(project_dir)
    for index, path in enumerate(files, start=1):
        rel = path.relative_to(project_dir)
        file_id = f"F{index:04d}"
        categories = classify(rel)
        category_codes = [code for code, _, _ in categories]
        labels = [label for _, label, _ in categories]
        sections = sorted({section for _, _, section in categories})
        suffix = path.suffix.lower()
        business_only = "business" in category_codes and len(category_codes) == 1
        deep_process = not business_only
        skip_reason = "商务资料只登记不深处理" if business_only else ""
        row = {
            "file_id": file_id,
            "relative_path": str(rel),
            "suffix": suffix or "无扩展名",
            "size_bytes": path.stat().st_size,
            "categories": "；".join(labels),
            "category_codes": category_codes,
            "deep_process": "yes" if deep_process else "no",
            "skip_reason": skip_reason,
            "applicable_sections": "；".join(sections),
        }
        file_rows.append(row)

        task = {
            "file_id": file_id,
            "relative_path": str(rel),
            "actions": [],
            "status": "skipped" if not deep_process else "processed",
            "notes": skip_reason,
        }

        if deep_process and suffix in TEXT_SUFFIXES | DOCX_SUFFIXES | PDF_SUFFIXES:
            text, issue, method = extract_text(path, args.max_chars)
            if text:
                text_path = text_dir / f"{file_id}.txt"
                text_path.write_text(text, encoding="utf-8")
                chunk_count = write_text_chunks(
                    text,
                    file_id=file_id,
                    source_file=str(rel),
                    text_dir=text_chunk_dir,
                    output_dir=output_dir,
                    index_path=processing_dir / "text_chunk_index.jsonl",
                )
                append_jsonl(
                    processing_dir / "text_index.jsonl",
                    {
                        "text_id": f"{file_id}-TXT001",
                        "file_id": file_id,
                        "source_file": str(rel),
                        "text_path": str(text_path.relative_to(output_dir)),
                        "char_count": len(text),
                        "extract_method": method,
                        "quality": "needs_review" if issue else "direct_text",
                        "chunk_count": chunk_count,
                    },
                )
                task["actions"].append("extract_text")
                task["actions"].append("chunk_text")
            if issue:
                append_jsonl(
                    processing_dir / "issues.jsonl",
                    {
                        "issue_id": f"I{file_id}",
                        "file_id": file_id,
                        "issue_type": "text_extraction",
                        "message": issue,
                        "user_visible": False,
                    },
                )

        if deep_process and suffix in SHEET_SUFFIXES | {".kml", ".ovkml"}:
            target = table_dir / f"{file_id}{suffix if suffix else '.txt'}"
            shutil.copy2(path, target)
            append_jsonl(
                processing_dir / "table_index.jsonl",
                {
                    "table_id": f"{file_id}-T001",
                    "file_id": file_id,
                    "source_file": str(rel),
                    "table_path": str(target.relative_to(output_dir)),
                    "table_type": "spatial" if suffix in {".kml", ".ovkml"} else "table",
                    "notes": "原文件复制，待后续结构化解析",
                },
            )
            task["actions"].append("index_table_or_spatial")
            if is_project_corner_coordinate_source(rel):
                coordinate_rows, coordinate_issue = extract_project_corner_coordinates(path, file_id=file_id, rel=rel)
                for coordinate_row in coordinate_rows:
                    append_jsonl(processing_dir / "project_corner_coordinates.jsonl", coordinate_row)
                if coordinate_rows:
                    task["actions"].append("extract_project_corner_coordinates")
                if coordinate_issue:
                    append_jsonl(
                        processing_dir / "issues.jsonl",
                        {
                            "issue_id": f"I{file_id}-COORD",
                            "file_id": file_id,
                            "issue_type": "project_corner_coordinate_extraction",
                            "message": coordinate_issue,
                            "user_visible": True,
                        },
                    )

        if deep_process and (suffix in FIGURE_SUFFIXES or "figure" in category_codes):
            target = figure_dir / f"{file_id}{suffix if suffix else path.suffix}"
            shutil.copy2(path, target)
            append_jsonl(
                processing_dir / "figure_index.jsonl",
                {
                    "figure_id": f"{file_id}-FIG001",
                    "file_id": file_id,
                    "source_file": str(rel),
                    "figure_path": str(target.relative_to(output_dir)),
                    "figure_type": "image" if suffix in FIGURE_SUFFIXES else "document_or_map",
                    "applicable_sections": "；".join(sections),
                    "caption_candidate": path.stem,
                    "notes": "图件索引，后续报告拼装需判断是否入正文",
                },
            )
            task["actions"].append("index_figure")

        if deep_process and is_supporting_file_candidate(rel):
            image_paths: list[str] = []
            conversion_issue = None
            if suffix in PDF_SUFFIXES:
                image_paths, conversion_issue = convert_supporting_pdf_to_images(
                    path,
                    file_id=file_id,
                    rel=rel,
                    image_dir=supporting_image_dir,
                    output_dir=output_dir,
                )
                task["actions"].append("convert_supporting_pdf_to_images")
            append_jsonl(
                processing_dir / "supporting_files.jsonl",
                {
                    "support_file_id": f"{file_id}-SUP001",
                    "file_id": file_id,
                    "source_file": str(rel),
                    "support_type": supporting_file_type(rel),
                    "image_paths": image_paths,
                    "notes": "第三章（支持性文件取得情况）专用白名单：仅立项、选址、核准文件或文物调查回函可入正文；其他文件一概不放。",
                },
            )
            task["actions"].append("index_supporting_file")
            if conversion_issue:
                append_jsonl(
                    processing_dir / "issues.jsonl",
                    {
                        "issue_id": f"I{file_id}-SUP",
                        "file_id": file_id,
                        "issue_type": "supporting_pdf_image_conversion",
                        "message": conversion_issue,
                        "user_visible": True,
                    },
                )

        append_jsonl(processing_dir / "processing_tasks.jsonl", task)

        if deep_process and "business" not in category_codes:
            evidence_id = f"E{len(evidence_rows)+1:04d}"
            evidence = {
                "evidence_id": evidence_id,
                "source_file_id": file_id,
                "source_file": str(rel),
                "source_location": "文件级索引",
                "evidence_type": "；".join(labels),
                "original_excerpt": "",
                "structured_value": "",
                "use_mode": "待确认",
                "applicable_section": "；".join(sections),
                "confidence": "待核验",
                "confirmation_reason": "资料处理阶段仅建文件级证据，事实抽取阶段需细化到页码、章节或片段。",
            }
            evidence_rows.append(evidence)
            append_jsonl(evidence_dir / "evidence_register.jsonl", evidence)
            for section in sections:
                append_jsonl(
                    evidence_dir / "section_evidence_index.jsonl",
                    {
                        "section_id": section,
                        "section_title": section,
                        "evidence_id": evidence_id,
                        "fact_id": "",
                        "recommended_use": "待事实抽取阶段判断",
                        "notes": "",
                    },
                )

    write_files_csv(processing_dir / "files.csv", file_rows)
    write_evidence_csv(evidence_dir / "evidence_register.csv", evidence_rows)
    for required in [
        "text_index.jsonl",
        "text_chunk_index.jsonl",
        "table_index.jsonl",
        "figure_index.jsonl",
        "supporting_files.jsonl",
        "project_corner_coordinates.jsonl",
        "external_sources.jsonl",
        "issues.jsonl",
        "user_blocking_gaps.jsonl",
    ]:
        (processing_dir / required).touch(exist_ok=True)

    gaps = blocking_gaps(file_rows)
    gap_path = processing_dir / "user_blocking_gaps.jsonl"
    if gaps:
        gap_path.write_text("", encoding="utf-8")
        for gap in gaps:
            append_jsonl(gap_path, gap)

    blocked = bool(gaps)
    module_state = {
        "project_name": args.project_name,
        "project_root": str(project_dir),
        "workspace_root": str(output_dir),
        "current_stage": "material_processing",
        "completed_stages": [] if blocked else ["material_processing"],
        "blocked": blocked,
        "blocking_file": str(gap_path.relative_to(output_dir)) if blocked else "",
        "created_at": started_at,
        "updated_at": now_iso(),
    }
    write_json(output_dir / "module_state.json", module_state)

    prompt_path = prompt_dir / ("blocked_user_input.md" if blocked else "next_prompt_fact_extraction.md")
    write_next_prompt(prompt_path, args.project_name, output_dir, blocked)

    manifest = {
        "project_name": args.project_name,
        "project_dir": str(project_dir),
        "output_dir": str(output_dir),
        "generated_at": now_iso(),
        "outputs": {
            "files_csv": "processing_output/files.csv",
            "processing_tasks": "processing_output/processing_tasks.jsonl",
            "text_index": "processing_output/text_index.jsonl",
            "text_chunk_index": "processing_output/text_chunk_index.jsonl",
            "table_index": "processing_output/table_index.jsonl",
            "figure_index": "processing_output/figure_index.jsonl",
            "supporting_files": "processing_output/supporting_files.jsonl",
            "project_corner_coordinates": "processing_output/project_corner_coordinates.jsonl",
            "external_sources": "processing_output/external_sources.jsonl",
            "issues": "processing_output/issues.jsonl",
            "user_blocking_gaps": "processing_output/user_blocking_gaps.jsonl",
            "evidence_register": "evidence/evidence_register.jsonl",
            "section_evidence_index": "evidence/section_evidence_index.jsonl",
        },
    }
    write_json(processing_dir / "manifest.json", manifest)

    done = {
        "module_name": "material_processing",
        "status": "blocked" if blocked else "completed",
        "started_at": started_at,
        "finished_at": now_iso(),
        "input_files": [str(project_dir)],
        "output_files": list(manifest["outputs"].values()) + ["module_state.json", str(prompt_path.relative_to(output_dir))],
        "blocking_gaps_count": len(gaps),
        "issues_count": sum(1 for _ in (processing_dir / "issues.jsonl").open(encoding="utf-8")),
        "next_prompt": str(prompt_path.relative_to(output_dir)),
        "notes": "资料处理阶段完成；正式报告需要后续事实抽取、分析判断和报告拼装阶段生成完整图文表报告。",
    }
    write_json(run_state_dir / "material_processing.module_done.json", done)

    print(f"资料处理完成：{output_dir}")
    print(f"文件数：{len(file_rows)}")
    print(f"阻断缺项：{len(gaps)}")
    print(f"下一阶段提示：{prompt_path}")


if __name__ == "__main__":
    main()
