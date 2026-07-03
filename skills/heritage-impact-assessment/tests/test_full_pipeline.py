from __future__ import annotations

import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


SKILL_DIR = Path(__file__).resolve().parents[1]


class FullPipelineTest(unittest.TestCase):
    def test_pipeline_reaches_full_report_outputs(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            project_dir = root / "sample_project"
            output_dir = root / "work"
            project_dir.mkdir()
            (project_dir / "项目说明.txt").write_text(
                "海南区觉海寺改扩建配套设施改造项目，建设单位为海南区觉海寺，"
                "建设内容包括山门、念佛堂、寮房、禅堂、鼓楼，项目面积7386平方米。",
                encoding="utf-8",
            )
            (project_dir / "文物调查报告.txt").write_text(
                "涉及桌子山秦长城东风农场七队长城3段，为自治区级文物保护单位。"
                "项目位于建设控制地带，距长城本体最近约192米，不涉及保护范围。",
                encoding="utf-8",
            )
            (project_dir / "项目用地范围.kml").write_text(
                "<kml><Placemark><name>项目用地范围</name></Placemark></kml>",
                encoding="utf-8",
            )

            commands = [
                [
                    "run_material_processing.py",
                    "--项目资料目录",
                    str(project_dir),
                    "--输出目录",
                    str(output_dir),
                    "--项目名称",
                    "海南区觉海寺改扩建配套设施改造项目",
                    "--覆盖",
                ],
                ["run_fact_extraction.py", "--工作目录", str(output_dir), "--覆盖"],
                ["run_analysis.py", "--工作目录", str(output_dir), "--覆盖"],
                ["run_report_assembly.py", "--workspace", str(output_dir), "--覆盖"],
            ]
            for command in commands:
                result = subprocess.run(
                    [sys.executable, str(SKILL_DIR / "scripts" / command[0]), *command[1:]],
                    text=True,
                    capture_output=True,
                )
                self.assertEqual(result.returncode, 0, result.stderr)

            clean_md = output_dir / "report_clean.md"
            self.assertTrue(clean_md.exists())
            self.assertTrue((output_dir / "report_with_evidence.md").exists())
            self.assertTrue((output_dir / "report_evidence_map.jsonl").exists())
            self.assertTrue((output_dir / "report_clean.docx").exists())
            self.assertTrue((output_dir / "run_state" / "report_assembly.module_done.json").exists())

            text = clean_md.read_text(encoding="utf-8")
            for heading in [
                "一、总则",
                "二、建设项目涉及文物概况",
                "三、建设项目规划概况",
                "四、项目用地范围与文物空间分布关系",
                "五、建设项目可能对文物造成的影响分析与评估",
                "六、减缓措施建议",
                "七、文物影响评估结论及建议",
                "八、支撑法律法规及文件",
            ]:
                self.assertIn(heading, text)

            doc = Document(output_dir / "report_clean.docx")
            paragraph_text = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("目录", paragraph_text)
            self.assertIn("附表一", paragraph_text)
            self.assertGreaterEqual(len(doc.tables), 3)
            self.assertGreaterEqual(len(doc.tables[-1].columns), 10)
            self.assertIn("项目", doc.tables[-1].cell(0, 0).text)
            self.assertIn("影响因子", doc.tables[-1].cell(0, 2).text)


if __name__ == "__main__":
    unittest.main()
