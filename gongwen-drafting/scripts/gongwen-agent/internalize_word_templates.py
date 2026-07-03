#!/usr/bin/env python3
"""Extract Word templates into AI-readable Markdown and JSON notes."""

from __future__ import annotations

import json
import re
from dataclasses import dataclass, asdict
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def find_root(start: Path) -> tuple[Path, str]:
    current = start.resolve()
    for candidate in [current, *current.parents]:
        if (candidate / "SKILL.md").exists() and (candidate / "assets/templates").exists():
            return candidate, "skill"
        if (candidate / "AGENTS.md").exists() and (candidate / "2.公文模板").exists():
            return candidate, "workspace"
    raise SystemExit("ERROR: run inside 公文撰写资料库 or the gongwen-drafting skill.")


ROOT, MODE = find_root(Path(__file__))
if MODE == "skill":
    OUTPUT_DIR = ROOT / "references/knowledge-index/模板内化"
    TEMPLATE_DIRS = [
        ROOT / "assets/fixed-format",
        ROOT / "assets/templates",
    ]
else:
    OUTPUT_DIR = ROOT / "0.资料索引" / "模板内化"
    TEMPLATE_DIRS = [
        ROOT / "1.政策法规与规范",
        ROOT / "2.公文模板",
    ]

PLACEHOLDER_RE = re.compile(
    r"(XX+|xx+|XXX+|xxx+|20XX|待补充|请填写|项目A|附件\d+|XXXXX+)"
)

ALIGNMENT_NAMES = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    WD_ALIGN_PARAGRAPH.DISTRIBUTE: "distribute",
}


@dataclass
class ParagraphInfo:
    index: int
    text: str
    style: str
    alignment: str
    first_line_indent_pt: float | None
    left_indent_pt: float | None
    right_indent_pt: float | None
    space_before_pt: float | None
    space_after_pt: float | None
    line_spacing: str | None
    fonts: list[str]
    has_bold: bool
    has_highlight: bool


@dataclass
class SectionInfo:
    index: int
    page_width_cm: float | None
    page_height_cm: float | None
    top_margin_cm: float | None
    bottom_margin_cm: float | None
    left_margin_cm: float | None
    right_margin_cm: float | None


def length_to_pt(value) -> float | None:
    if value is None:
        return None
    return round(value.pt, 2)


def length_to_cm(value) -> float | None:
    if value is None:
        return None
    return round(value.cm, 2)


def alignment_name(value) -> str:
    if value is None:
        return "inherit"
    return ALIGNMENT_NAMES.get(value, str(value))


def safe_name(path: Path) -> str:
    name = path.stem
    name = re.sub(r"[\\/:\*\?\"<>\|]", "_", name)
    return name[:140]


def iter_docx_files() -> Iterable[Path]:
    seen: set[Path] = set()
    for base in TEMPLATE_DIRS:
        if not base.exists():
            continue
        for path in sorted(base.rglob("*.docx")):
            if path.name.startswith(("~$", ".~")):
                continue
            if path in seen:
                continue
            seen.add(path)
            yield path


def para_text(paragraph) -> str:
    return paragraph.text.replace("\u3000", " ").strip()


def paragraph_info(index: int, paragraph) -> ParagraphInfo:
    pformat = paragraph.paragraph_format
    fonts: set[str] = set()
    has_bold = False
    has_highlight = False
    for run in paragraph.runs:
        if run.font.name:
            fonts.add(run.font.name)
        if run.bold:
            has_bold = True
        try:
            if run.font.highlight_color:
                has_highlight = True
        except ValueError:
            # Some Word files store highlight="none", which python-docx cannot map.
            pass
    return ParagraphInfo(
        index=index,
        text=para_text(paragraph),
        style=paragraph.style.name if paragraph.style else "",
        alignment=alignment_name(paragraph.alignment),
        first_line_indent_pt=length_to_pt(pformat.first_line_indent),
        left_indent_pt=length_to_pt(pformat.left_indent),
        right_indent_pt=length_to_pt(pformat.right_indent),
        space_before_pt=length_to_pt(pformat.space_before),
        space_after_pt=length_to_pt(pformat.space_after),
        line_spacing=str(pformat.line_spacing) if pformat.line_spacing else None,
        fonts=sorted(fonts),
        has_bold=has_bold,
        has_highlight=has_highlight,
    )


def section_info(index: int, section) -> SectionInfo:
    return SectionInfo(
        index=index,
        page_width_cm=length_to_cm(section.page_width),
        page_height_cm=length_to_cm(section.page_height),
        top_margin_cm=length_to_cm(section.top_margin),
        bottom_margin_cm=length_to_cm(section.bottom_margin),
        left_margin_cm=length_to_cm(section.left_margin),
        right_margin_cm=length_to_cm(section.right_margin),
    )


def table_texts(document: Document) -> list[list[list[str]]]:
    tables: list[list[list[str]]] = []
    for table in document.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)
    return tables


def infer_doc_type(name: str) -> str:
    checks = [
        ("考古勘探工作支持请示", ["考古勘探工作支持", "协助勘探"]),
        ("文物保护许可申请", ["文物保护许可", "有关文物事宜", "文物核查", "文物调查事宜", "办理文物调查"]),
        ("申请开展考古勘探工作请示", ["申请开展", "考古勘探工作"]),
        ("勘探计划备案请示", ["勘探计划备案", "工作计划"]),
        ("勘探报告备案请示", ["勘探成果报告备案", "勘探报告备案"]),
        ("勘探验收请示", ["勘探验收", "勘查验收许可"]),
        ("开工请示", ["开工"]),
        ("发掘请示", ["发掘"]),
        ("文物保护安全责任书", ["安全责任书"]),
        ("文物调查报告", ["文物调查的报告", "文物调查报告", "市级调查报告"]),
        ("承诺书", ["承诺书"]),
    ]
    for doc_type, words in checks:
        if any(word in name for word in words):
            return doc_type
    return "其他"


def markdown_for_template(path: Path, payload: dict) -> str:
    rel = path.relative_to(ROOT)
    paragraphs: list[dict] = payload["paragraphs"]
    body_lines = [p["text"] for p in paragraphs if p["text"]]
    placeholders = sorted(set(payload["placeholders"]))
    highlighted = [p for p in paragraphs if p["has_highlight"] and p["text"]]
    sections = payload["sections"]
    tables = payload["tables"]

    lines: list[str] = []
    lines.append(f"# {path.stem}")
    lines.append("")
    lines.append("## 来源")
    lines.append("")
    lines.append(f"- 原文件：[[{rel.as_posix()}]]")
    lines.append(f"- 内化时间：{payload['generated_at']}")
    lines.append(f"- 推断文种：{payload['doc_type']}")
    lines.append(f"- 段落数：{len(paragraphs)}")
    lines.append(f"- 表格数：{len(tables)}")
    lines.append("")
    lines.append("## 版式信息")
    lines.append("")
    if sections:
        lines.append("| 节 | 纸宽cm | 纸高cm | 上cm | 下cm | 左cm | 右cm |")
        lines.append("| ---: | ---: | ---: | ---: | ---: | ---: | ---: |")
        for s in sections:
            lines.append(
                "| {index} | {page_width_cm} | {page_height_cm} | {top_margin_cm} | "
                "{bottom_margin_cm} | {left_margin_cm} | {right_margin_cm} |".format(**s)
            )
    else:
        lines.append("- 未提取到节信息。")
    lines.append("")
    lines.append("## 占位与高亮")
    lines.append("")
    if placeholders:
        lines.append("- 占位词：" + "、".join(placeholders))
    else:
        lines.append("- 占位词：未发现常见占位词。")
    if highlighted:
        lines.append("- 含高亮段落：")
        for p in highlighted[:20]:
            lines.append(f"  - P{p['index']}: {p['text']}")
    else:
        lines.append("- 含高亮段落：未发现。")
    lines.append("")
    lines.append("## 正文骨架")
    lines.append("")
    if body_lines:
        for text in body_lines:
            lines.append(text)
            lines.append("")
    else:
        lines.append("- 未提取到正文文本。")
        lines.append("")
    if tables:
        lines.append("## 表格文本")
        lines.append("")
        for i, table in enumerate(tables, 1):
            lines.append(f"### 表格 {i}")
            lines.append("")
            for row in table:
                joined = " | ".join(cell.replace("\n", " / ") for cell in row)
                lines.append(f"- {joined}")
            lines.append("")
    lines.append("## 段落版式明细")
    lines.append("")
    lines.append(
        "| P | 样式 | 对齐 | 首行pt | 左缩进pt | 右缩进pt | 字体 | 加粗 | 高亮 | 文本 |"
    )
    lines.append("| ---: | --- | --- | ---: | ---: | ---: | --- | --- | --- | --- |")
    for p in paragraphs:
        text = p["text"].replace("|", "\\|")
        if len(text) > 80:
            text = text[:77] + "..."
        fonts = ",".join(p["fonts"])
        lines.append(
            f"| {p['index']} | {p['style']} | {p['alignment']} | "
            f"{p['first_line_indent_pt']} | {p['left_indent_pt']} | {p['right_indent_pt']} | "
            f"{fonts} | {p['has_bold']} | {p['has_highlight']} | {text} |"
        )
    lines.append("")
    return "\n".join(lines)


def extract_template(path: Path, generated_at: str) -> dict:
    document = Document(path)
    paragraphs = [asdict(paragraph_info(i, p)) for i, p in enumerate(document.paragraphs, 1)]
    sections = [asdict(section_info(i, s)) for i, s in enumerate(document.sections, 1)]
    tables = table_texts(document)
    all_text = "\n".join([p["text"] for p in paragraphs])
    for table in tables:
        for row in table:
            all_text += "\n" + "\n".join(row)
    placeholders = PLACEHOLDER_RE.findall(all_text)
    return {
        "source": path.relative_to(ROOT).as_posix(),
        "name": path.stem,
        "doc_type": infer_doc_type(path.stem),
        "generated_at": generated_at,
        "sections": sections,
        "paragraphs": paragraphs,
        "tables": tables,
        "placeholders": placeholders,
    }


def write_outputs() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    manifest = []
    for path in iter_docx_files():
        payload = extract_template(path, generated_at)
        md_name = f"{safe_name(path)}.md"
        json_name = f"{safe_name(path)}.json"
        md_path = OUTPUT_DIR / md_name
        json_path = OUTPUT_DIR / json_name
        md_path.write_text(markdown_for_template(path, payload), encoding="utf-8")
        json_path.write_text(
            json.dumps(payload, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        manifest.append(
            {
                "source": payload["source"],
                "name": payload["name"],
                "doc_type": payload["doc_type"],
                "markdown": f"0.资料索引/模板内化/{md_name}",
                "json": f"0.资料索引/模板内化/{json_name}",
                "paragraph_count": len(payload["paragraphs"]),
                "table_count": len(payload["tables"]),
            }
        )
    write_manifest(manifest, generated_at)


def write_manifest(manifest: list[dict], generated_at: str) -> None:
    json_path = OUTPUT_DIR / "manifest.json"
    json_path.write_text(
        json.dumps(
            {
                "generated_at": generated_at,
                "template_count": len(manifest),
                "templates": manifest,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    lines = [
        "# 模板内化总览",
        "",
        f"- 内化时间：{generated_at}",
        f"- 模板数量：{len(manifest)}",
        "",
        "## 使用方式",
        "",
        "- 写作前先读对应文种作战手册，再读取本目录中对应模板内化文件。",
        "- Markdown 记录正文骨架和版式线索，JSON 保留更完整的段落和表格数据。",
        "- 原始 Word 模板不在此处修改；如模板变化，重新运行 `python3 agent/internalize_word_templates.py`。",
        "",
        "## 模板清单",
        "",
        "| 文种 | 模板 | 内化 Markdown | 原文件 |",
        "| --- | --- | --- | --- |",
    ]
    for item in manifest:
        lines.append(
            f"| {item['doc_type']} | {item['name']} | "
            f"[[{item['markdown']}|Markdown]] | [[{item['source']}|原文件]] |"
        )
    lines.append("")
    (OUTPUT_DIR / "00-模板内化总览.md").write_text("\n".join(lines), encoding="utf-8")


if __name__ == "__main__":
    write_outputs()
