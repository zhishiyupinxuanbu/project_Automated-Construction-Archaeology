#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""固定公文版式工具。

本模块把高频公文版式规则落到代码里，减少跨电脑、跨会话时只靠
自然语言规则导致的漂移。重点服务文物保护许可申请/核查请示。
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, List, Sequence
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BODY_FONT = "仿宋_GB2312"
TITLE_FONT = "方正小标宋_GBK"
BODY_SIZE = 16
TITLE_SIZE = 22
RED_HEADER_SIZE = 36
LINE_SPACING = 28
TITLE_LINE_SPACING = 30
RED_LINE_BORDER_SIZE = 24
RED_COLOR = RGBColor(0xFF, 0x00, 0x00)
TWO_CHARS = Pt(BODY_SIZE * 2)
FOUR_CHARS = Pt(BODY_SIZE * 4)
DOC_NUMBER_PLACEHOLDER = "【文号】"
SIGNER_PLACEHOLDER = "【签发人】"

# 附件区域独立于正文首行缩进：
# 附件：从左空二字处起排；第 2 项起序号与第 1 项的“1.”对齐；
# 附件内容和续行统一对齐。
ATTACH_LABEL_START = Pt(BODY_SIZE * 2)
ATTACH_NUMBER_START = Pt(BODY_SIZE * 5)
ATTACH_CONTENT_START = Pt(BODY_SIZE * 7)

NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def text_width_pt(text: str, size: int = BODY_SIZE) -> float:
    width = 0.0
    for char in text or "":
        width += size * 0.55 if char.isascii() else size
    return width


def signature_center_position(date_text: str) -> Pt:
    content_width = Cm(21 - 2.8 - 2.6)
    date_right_edge_pt = content_width.pt - FOUR_CHARS.pt
    return Pt(max(0, date_right_edge_pt - text_width_pt(date_text) / 2))


def red_header_font_size(text: str) -> float:
    content_width_pt = Cm(21 - 2.8 - 2.6).pt
    unit_width = text_width_pt(text, size=1)
    if unit_width <= 0:
        return RED_HEADER_SIZE
    return max(22, min(RED_HEADER_SIZE, content_width_pt / unit_width))


@dataclass
class HechaRequestData:
    issuing_org: str
    recipient_org: str
    project_name: str
    location: str = ""
    scale: str = ""
    project_area: str = ""
    approval_file: str = "审批机关立项选址核准文件"
    basis_sentence: str = "按照各级文物管理要求及《内蒙古自治区文物局关于做好基本建设用地考古工作的通知》（内文物发〔2025〕6号）文件要求"
    date_text: str = ""
    doc_number: str = ""
    signer: str = ""
    add_red_header: bool = True
    body_paragraphs: List[str] = field(default_factory=list)
    attachments: List[str] = field(default_factory=list)


def set_run_font(
    run,
    font_name: str = BODY_FONT,
    size: float = BODY_SIZE,
    *,
    color: RGBColor | None = None,
    bold: bool | None = None,
) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.font.bold = bold


def set_paragraph_bottom_border(
    paragraph,
    *,
    color: str = "FF0000",
    size: int = RED_LINE_BORDER_SIZE,
) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def apply_document_geometry(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    section.bottom_margin = Cm(3.5)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(BODY_SIZE)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = Pt(LINE_SPACING)


def add_red_header(
    doc: Document,
    issuing_org: str,
    *,
    doc_number: str = "",
    signer: str = "",
) -> None:
    """Add the project default red letterhead before the title."""
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_before = Pt(0)
    header.paragraph_format.space_after = Pt(0)
    header.paragraph_format.line_spacing = 1
    header_text = f"{issuing_org}文件"
    header_run = header.add_run(header_text)
    set_run_font(
        header_run,
        TITLE_FONT,
        red_header_font_size(header_text),
        color=RED_COLOR,
    )

    info = doc.add_paragraph()
    info.paragraph_format.space_before = Pt(0)
    info.paragraph_format.space_after = Pt(0)
    info.paragraph_format.line_spacing = Pt(LINE_SPACING)
    info.paragraph_format.first_line_indent = TWO_CHARS
    info.paragraph_format.tab_stops.clear_all()
    info.paragraph_format.tab_stops.add_tab_stop(
        Cm(21 - 2.8 - 2.6),
        WD_TAB_ALIGNMENT.RIGHT,
    )
    number_run = info.add_run(doc_number.strip() if doc_number.strip() else DOC_NUMBER_PLACEHOLDER)
    set_run_font(number_run)
    info.add_run("\t")
    signer_text = signer.strip() if signer.strip() else SIGNER_PLACEHOLDER
    signer_run = info.add_run(f"签发人：{signer_text}")
    set_run_font(signer_run)

    line = doc.add_paragraph()
    line.paragraph_format.space_before = Pt(0)
    line.paragraph_format.space_after = Pt(0)
    line.paragraph_format.line_spacing = Pt(1)
    set_paragraph_bottom_border(line)

    blank = doc.add_paragraph()
    blank.alignment = WD_ALIGN_PARAGRAPH.CENTER
    blank.paragraph_format.space_before = Pt(0)
    blank.paragraph_format.space_after = Pt(0)
    blank.paragraph_format.line_spacing = Pt(LINE_SPACING)


def add_title(doc: Document, lines: Sequence[str] | str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(24)
    paragraph.paragraph_format.line_spacing = Pt(TITLE_LINE_SPACING)
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.right_indent = None
    title = lines if isinstance(lines, str) else "".join(lines)
    run = paragraph.add_run(title)
    set_run_font(run, TITLE_FONT, TITLE_SIZE)


def add_body_paragraph(doc: Document, text: str, *, first_line: bool = True) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(LINE_SPACING)
    if first_line:
        paragraph.paragraph_format.first_line_indent = TWO_CHARS
    run = paragraph.add_run(text)
    set_run_font(run)


def add_attachment_item(doc: Document, index: int, text: str) -> None:
    paragraph = doc.add_paragraph()
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(LINE_SPACING)
    pf.left_indent = ATTACH_CONTENT_START
    if index == 1:
        pf.first_line_indent = ATTACH_LABEL_START - ATTACH_CONTENT_START
        line = f"附件：1.\t{text}"
    else:
        pf.first_line_indent = ATTACH_NUMBER_START - ATTACH_CONTENT_START
        line = f"{index}.\t{text}"
    pf.tab_stops.clear_all()
    pf.tab_stops.add_tab_stop(ATTACH_CONTENT_START)
    run = paragraph.add_run(line)
    set_run_font(run)


def add_signature(doc: Document, issuing_org: str, date_text: str) -> None:
    doc.add_paragraph()

    # 单一机关行文：先基于日期居中。只有居中后换行/不稳时，才可另行靠右。
    signature = doc.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.LEFT
    signature.paragraph_format.left_indent = Pt(0)
    signature.paragraph_format.right_indent = Pt(0)
    signature.paragraph_format.space_after = Pt(0)
    signature.paragraph_format.line_spacing = Pt(LINE_SPACING)
    signature.paragraph_format.tab_stops.clear_all()
    signature.paragraph_format.tab_stops.add_tab_stop(
        signature_center_position(date_text),
        WD_TAB_ALIGNMENT.CENTER,
    )
    run = signature.add_run("\t" + issuing_org)
    set_run_font(run)

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date.paragraph_format.right_indent = FOUR_CHARS
    date.paragraph_format.space_after = Pt(0)
    date.paragraph_format.line_spacing = Pt(LINE_SPACING)
    run = date.add_run(date_text)
    set_run_font(run)


def default_hecha_attachments(data: HechaRequestData) -> List[str]:
    project = data.project_name
    first = data.approval_file or "审批机关立项选址核准文件"
    return [
        first,
        "企业法人营业执照",
        "法定代表人身份证复印件",
        f"{project}用地经纬度坐标 Excel 表",
        f"{project}用地范围 KML 格式坐标文件",
        f"{project}用地宗地图",
    ]


def default_hecha_body(data: HechaRequestData) -> List[str]:
    facts = [f"{data.issuing_org}{data.project_name}"]
    if data.approval_file:
        facts.append(f"已经《{data.approval_file}》批准或核准")
    if data.location:
        facts.append(f"项目位于{data.location}")
    if data.project_area:
        facts.append(f"项目面积{data.project_area}")
    if data.scale:
        facts.append(f"建设规模及主要建设内容为{data.scale}")
    intro = "，".join(facts) + "。为依法依规办理项目用地范围内文物保护相关手续，现就该项目用地范围内文物保护许可事宜请示如下。"
    request = (
        f"{data.basis_sentence}，现申请贵局对该项目用地范围内是否涉及文物保护单位、"
        "不可移动文物及其保护范围、建设控制地带等情况进行核查，并请出具相关审查意见。"
    )
    return [intro, request, "妥否，请批示。"]


def create_hecha_docx(data: HechaRequestData, output_path: Path) -> Path:
    doc = Document()
    apply_document_geometry(doc)

    if data.add_red_header:
        add_red_header(
            doc,
            data.issuing_org,
            doc_number=data.doc_number,
            signer=data.signer,
        )
    add_title(doc, f"{data.issuing_org}关于办理{data.project_name}用地范围内文物保护许可的请示")
    add_body_paragraph(doc, f"{data.recipient_org}：", first_line=False)
    for text in data.body_paragraphs or default_hecha_body(data):
        add_body_paragraph(doc, text)

    doc.add_paragraph()
    attachments = data.attachments or default_hecha_attachments(data)
    for index, item in enumerate(attachments, start=1):
        add_attachment_item(doc, index, item)

    add_signature(doc, data.issuing_org, data.date_text)
    doc.save(output_path)
    return output_path


def collect_text(docx_path: Path) -> str:
    return "\n".join(paragraph.text for paragraph in Document(docx_path).paragraphs)


def _paragraph_xml_roots(docx_path: Path) -> List[ET.Element]:
    with ZipFile(docx_path) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    return root.findall(".//w:p", NS)


def _paragraph_has_red_bottom_border(paragraph) -> bool:
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return False
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        return False
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        return False
    color = bottom.get(qn("w:color"), "").upper()
    val = bottom.get(qn("w:val"), "")
    return val not in {"", "nil", "none"} and color in {"FF0000", "FF0000".lower()}


def _red_header_info_errors(paragraphs) -> List[str]:
    info_para = next((p for p in paragraphs[:5] if "签发人" in p.text), None)
    if info_para is None:
        return ["红头缺少文号/签发人行"]

    info_text = " ".join(info_para.text.replace("\t", " ").split())
    before_signer = info_text.split("签发人", 1)[0].strip()
    errors: List[str] = []
    if before_signer in {"", "文号"}:
        errors.append("红头文号缺少“【文号】”占位或正式文号")
    if SIGNER_PLACEHOLDER not in info_text and info_text.rstrip() in {"签发人：", "签发人:"}:
        errors.append("红头签发人缺少“【签发人】”占位或正式姓名")
    if before_signer and before_signer != "文号" and DOC_NUMBER_PLACEHOLDER not in before_signer and "号" not in before_signer:
        errors.append("红头文号不是“【文号】”占位，也未识别到正式文号")
    after_signer = info_text.split("签发人", 1)[1].lstrip("：: ").strip()
    if not after_signer:
        errors.append("红头签发人缺少“【签发人】”占位或正式姓名")
    return errors


def audit_hecha_docx(docx_path: Path) -> List[str]:
    errors: List[str] = []
    doc = Document(docx_path)
    if not doc.paragraphs or doc.paragraphs[0].text.strip() == "":
        errors.append("文档开头缺少红色发文机关标志")
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    if not paragraphs:
        return ["文档为空"]

    text = collect_text(docx_path)
    is_wulanchabu_template = (
        ("乌兰察布" in text or "察哈尔右翼后旗" in text)
        and any("文物核查的请示" in p.text for p in paragraphs[:3])
    )
    is_kulun_template = "库伦旗" in text and any(
        "进行文物核查的请示" in p.text.replace("\n", "")
        for p in paragraphs[:8]
    )

    has_red_header = paragraphs[0].text.endswith("文件")
    if not has_red_header:
        errors.append("缺少红头：第一段应为发文机关标志，如“发文单位文件”")
    else:
        first_run = next((run for run in paragraphs[0].runs if run.text.strip()), None)
        color = first_run.font.color.rgb if first_run and first_run.font.color else None
        if color != RED_COLOR:
            errors.append("红头发文机关标志不是红色")

    red_line = next((p for p in doc.paragraphs[:6] if _paragraph_has_red_bottom_border(p)), None)
    if red_line is None:
        errors.append("缺少红色分隔线，或红线使用了字符拼接而非连续段落边框")
    if has_red_header:
        errors.extend(_red_header_info_errors(paragraphs))

    if is_wulanchabu_template:
        title = next((p for p in paragraphs[:7] if "文物核查的请示" in p.text), paragraphs[0])
    elif is_kulun_template:
        title = next(
            (
                p
                for p in paragraphs[:8]
                if "进行文物核查的请示" in p.text.replace("\n", "")
            ),
            paragraphs[0],
        )
    else:
        title = next((p for p in paragraphs[:7] if "关于办理" in p.text), paragraphs[0])

    if (
        not (is_wulanchabu_template or is_kulun_template)
        and "关于办理" not in paragraphs[0].text
        and title is paragraphs[0]
    ):
        errors.append("未识别到核查请示标题")

    title_index = next((i for i, p in enumerate(doc.paragraphs) if p is title), 0)
    title_xml = _paragraph_xml_roots(docx_path)[title_index]
    if (
        not (is_wulanchabu_template or is_kulun_template)
        and title_xml.find(".//w:br", NS) is not None
    ):
        errors.append("标题不应手动换行，应为一长句居中")
    for run in title.runs:
        if run.text.strip():
            if (
                run.font.size
                and not (is_wulanchabu_template or is_kulun_template)
                and abs(run.font.size.pt - TITLE_SIZE) > 0.1
            ):
                errors.append("标题字号不是 2 号小标宋对应的 22pt")
            east_asia = run._element.rPr.rFonts.get(qn("w:eastAsia")) if run._element.rPr is not None else ""
            if not (is_wulanchabu_template or is_kulun_template) and east_asia != TITLE_FONT:
                errors.append(f"标题字体不是 {TITLE_FONT}")

    bad_terms = ["调查面积", "勘探面积", "待补充", "20XX", "XX"]
    for term in bad_terms:
        if term in text:
            errors.append(f"正文含禁用或占位表述：{term}")
    for term in ["考古调查勘探报告", "文物保护方案", "文物影响评估报告"]:
        if term in text:
            errors.append(f"核查请示附件不应默认包含：{term}")

    attachment_prefixes = tuple(f"{index}." for index in range(2, 21))
    attachment_paras = [
        p for p in paragraphs
        if p.text.startswith("附件：") or p.text.startswith(attachment_prefixes)
    ]
    if attachment_paras:
        if not attachment_paras[0].text.startswith("附件：1."):
            errors.append("附件第一项没有连续起排为“附件：1.”")
        expected = ["附件：1."] + [f"{index}." for index in range(2, len(attachment_paras) + 1)]
        for prefix, para in zip(expected, attachment_paras):
            if not para.text.startswith(prefix):
                errors.append(f"附件序号异常：应以 {prefix} 开头")
    else:
        errors.append("未识别到附件说明")

    for para in paragraphs:
        text = para.text.strip()
        if not text or text == paragraphs[0].text.strip():
            continue
        if text.endswith("：") or text.startswith("附件：") or text.startswith(attachment_prefixes):
            continue
        if text in {"妥否，请批示。"} or len(text) > 25:
            indent = para.paragraph_format.first_line_indent
            if indent is not None and abs(indent.pt - (BODY_SIZE * 2)) > 0.1:
                errors.append(f"正文段落首行缩进不是两个中文字宽：{text[:16]}")

    signature_para = next((p for p in reversed(paragraphs) if p.text.strip() and p.text.strip() == paragraphs[-2].text.strip()), None)
    if signature_para is not None:
        if len(list(signature_para.paragraph_format.tab_stops)) == 0:
            errors.append("发文单位落款未使用居中制表位对准日期中心")

    return errors


def write_audit_report(docx_path: Path, report_path: Path) -> Path:
    errors = audit_hecha_docx(docx_path)
    lines = [f"# {docx_path.name} 版式校验", ""]
    if errors:
        lines.append("校验未通过：")
        lines.extend(f"- {item}" for item in errors)
    else:
        lines.append("校验通过。")
    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return report_path
