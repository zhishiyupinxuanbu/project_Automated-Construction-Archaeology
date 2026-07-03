#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
本地公文撰写 Agent MVP。

功能：
- 扫描资料库，解析 DOCX 与可复制文字 PDF，建立本地索引。
- 基于事项类型、旗区、项目关键词检索模板和参考件。
- 根据用户输入生成正文和 Word 成稿，使用资料库内固定版式参数控制关键格式。
- 做关键信息一致性校验，降低模型式改写风险。
- 提供命令行和零依赖本地 Web 表单。
"""

from __future__ import annotations

import argparse
import datetime as dt
import hashlib
import html
import json
import os
import re
import shutil
import subprocess
import sys
import textwrap
import tempfile
import urllib.parse
from collections import Counter
from dataclasses import asdict, dataclass, field
from http import HTTPStatus
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple
from xml.etree import ElementTree as ET
from zipfile import BadZipFile, ZipFile

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except Exception:  # pragma: no cover - handled at runtime
    Document = None
    WD_ALIGN_PARAGRAPH = None
    WD_TAB_ALIGNMENT = None
    OxmlElement = None
    qn = None
    Cm = None
    Pt = None


ROOT = Path(__file__).resolve().parents[1]
WORK_DIR_NAME = ".gongwen_agent"
INDEX_NAME = "index.json"
TEXT_DIR_NAME = "texts"
OUTPUT_DIR_NAME = "outputs"

SUPPORTED_EXTS = {".docx", ".pdf"}
IGNORED_PARTS = {".temp", "__macosx"}
MIN_USEFUL_TEXT = 80
DEFAULT_TOP_K = 5
DEFAULT_OCR_DPI = 220


def default_output_dir(base_dir: Path = ROOT) -> Path:
    desktop = Path.home() / "Desktop"
    if desktop.exists():
        return desktop
    return base_dir / WORK_DIR_NAME / OUTPUT_DIR_NAME


def result_path_for_display(path: Path, base_dir: Path = ROOT) -> str:
    try:
        return str(path.relative_to(base_dir))
    except ValueError:
        return str(path)

BUSINESS_ALIASES = {
    "政策法规与规范": ["政策法规", "公文格式", "格式范本", "固定版式"],
    "文物保护许可申请": ["文物保护许可", "保护许可", "有关文物事宜", "文物事宜", "文物核查"],
    "申请开展考古勘探工作请示": ["勘探请示", "开展勘探请示", "申请勘探", "申请开展考古勘探", "考古勘探工作请示"],
    "勘探验收请示": ["勘探验收", "勘查验收", "验收许可", "勘探工作验收"],
    "发掘请示": ["发掘", "抢救性考古发掘", "遗迹进行考古发掘"],
    "开工请示": ["开工", "验收后开工", "开工请示"],
    "勘探计划备案请示": ["勘探计划备案", "计划备案"],
    "勘探报告备案请示": ["勘探成果报告备案", "报告备案", "成果报告备案"],
    "文物保护安全责任书": ["安全责任书", "承诺书", "风险评估", "安全防护方案"],
    "文物调查意见": ["文物审查意见", "调查意见", "审查意见"],
    "其他": ["说明", "情况说明"],
    "文物调查报告": ["文物调查报告", "调查报告"],
}

REGION_ALIASES = {
    "伊旗": "伊金霍洛旗",
    "伊金霍洛旗": "伊金霍洛旗",
    "准旗": "准格尔旗",
    "准格尔旗": "准格尔旗",
    "乌审旗": "乌审旗",
    "杭锦旗": "杭锦旗",
    "达旗": "达拉特旗",
    "达拉特旗": "达拉特旗",
    "鄂托克旗": "鄂托克旗",
    "呼伦贝尔": "呼伦贝尔",
}

DEFAULT_ATTACHMENTS = {
    "文物保护许可申请": [
        "审批机关立项选址核准文件",
        "企业法人营业执照",
        "法定代表人身份证复印件",
        "XX项目用地经纬度坐标 Excel 表",
        "XX项目用地范围 KML 格式坐标文件",
        "XX项目用地宗地图",
    ],
    "勘探验收请示": [
        "XX项目考古调查、勘探工作报告",
    ],
    "申请开展考古勘探工作请示": [],
    "发掘请示": [
        "考古发掘工作方案",
        "文物遗迹发现和勘探材料",
        "项目用地范围及位置图",
        "建设单位申请文件",
    ],
    "开工请示": [
        "文物勘探验收意见或相关批复",
        "建设单位开工申请文件",
        "项目用地范围材料",
    ],
    "勘探计划备案请示": [
        "文物勘探计划",
        "项目基本情况说明",
        "用地范围坐标资料",
    ],
    "勘探报告备案请示": [
        "文物勘探成果报告",
        "项目用地范围材料",
        "建设单位申请或说明文件",
    ],
    "文物调查意见": [
        "项目选址范围资料",
        "文物调查或核查过程材料",
        "建设单位申请文件",
    ],
    "文物调查报告": [
        "基本建设项目文物调查表",
        "拟建项目用地范围图及相关示意图",
        "现场照片",
        "拟建项目用地坐标",
        "项目特殊情况说明",
    ],
}


@dataclass
class FileRecord:
    id: str
    path: str
    extension: str
    business_type: str
    region: str
    source_role: str
    is_template: bool
    title: str
    project_name_guess: str
    issuing_org_guess: str
    date_guess: str
    text_chars: int
    parse_status: str
    text_path: str


@dataclass
class ProjectInput:
    business_type: str = ""
    issuing_org: str = ""
    recipient_org: str = ""
    project_name: str = ""
    construction_unit: str = ""
    location: str = ""
    scale: str = ""
    land_area: str = ""
    approval_basis: str = ""
    contact: str = ""
    phone: str = ""
    special_notes: str = ""
    region: str = ""

    @classmethod
    def from_mapping(cls, data: Dict[str, str]) -> "ProjectInput":
        values = {field_name: clean_text(str(data.get(field_name, ""))) for field_name in cls.__dataclass_fields__}
        if not values.get("region"):
            values["region"] = infer_region(" ".join(values.values()))
        return cls(**values)


@dataclass
class DraftResult:
    classified_business_type: str
    selected_template: Optional[FileRecord]
    references: List[FileRecord]
    docx_path: str
    checks: List[Dict[str, str]]
    markdown_path: str = ""
    warnings: List[str] = field(default_factory=list)


def clean_text(text: Optional[str]) -> str:
    if text is None:
        return ""
    text = text.replace("\u3000", " ").replace("\xa0", " ")
    text = re.sub(r"[ \t\r\f\v]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def safe_filename(name: str, suffix: str = "") -> str:
    name = re.sub(r"[\\/:*?\"<>|]+", "_", name).strip()
    name = re.sub(r"\s+", " ", name)
    if not name:
        name = "未命名公文"
    return name[:120] + suffix


def stable_id(path: Path) -> str:
    return hashlib.sha1(str(path.relative_to(ROOT)).encode("utf-8")).hexdigest()[:16]


def strip_business_prefix(part: str) -> str:
    return re.sub(r"^\d+\.", "", part).strip()


def infer_business_from_path(path: Path) -> str:
    try:
        parts = path.relative_to(ROOT).parts
    except Exception:
        parts = path.parts

    if any("政策法规与规范" in part for part in parts):
        return "政策法规与规范"

    generic_dirs = {"公文参考", "公文模板", "政策法规与规范"}
    candidates = [strip_business_prefix(part) for part in parts]
    for candidate in candidates:
        if candidate in BUSINESS_ALIASES:
            return candidate

    joined = " ".join(candidates + [path.stem])
    for business_type, aliases in BUSINESS_ALIASES.items():
        if business_type in joined or any(alias in joined for alias in aliases):
            return business_type

    for candidate in candidates:
        if candidate and candidate not in generic_dirs:
            return candidate
    return strip_business_prefix(parts[0]) if parts else ""


def infer_region(text: str) -> str:
    for alias, canonical in REGION_ALIASES.items():
        if alias in text:
            return canonical
    return ""


def infer_source_role(path: Path) -> str:
    parts = [p.lower() for p in path.parts]
    joined = "/".join(parts)
    if "企业申请文件" in joined:
        return "企业申请文件"
    if "旗区给市局申请文件" in joined:
        return "旗区给市局申请文件"
    if "参考文件" in joined:
        return "参考文件"
    if "模版" in joined or "模板" in joined:
        return "模板"
    return "样例"


def is_template_file(path: Path) -> bool:
    name = path.name
    return bool(re.search(r"(模版|模板|通用模版|XXX|XX|xxx)", name))


def title_from_name(path: Path) -> str:
    title = path.stem
    title = re.sub(r"^\d{6,8}\s*", "", title)
    title = re.sub(r"^(通用)?模[版板][—-]*", "", title)
    return clean_text(title)


def guess_project_name(title: str) -> str:
    text = title
    if "关于" in text:
        text = text.split("关于", 1)[1]
    text = re.sub(r"(用地范围|选址范围|建设范围|项目范围).*$", "", text)
    text = re.sub(r"(有关文物.*|文物保护许可.*|文物核查.*|文物审查.*|文物勘探.*|文物调查.*|进行文物.*|申请办理.*|申请.*)$", "", text)
    text = re.sub(r"(的请示|的函|的报告|的说明)$", "", text)
    return clean_text(text)


def guess_issuing_org(title: str) -> str:
    if "关于" in title:
        candidate = title.split("关于", 1)[0]
        candidate = re.sub(r"^(模板|模版|通用模版)[—-]*", "", candidate)
        return clean_text(candidate)
    return ""


def guess_date(text: str, path: Path) -> str:
    def valid_date(value: str) -> str:
        try:
            dt.date.fromisoformat(value)
        except ValueError:
            return ""
        return value

    joined = path.name + "\n" + text[:2000]
    match = re.search(r"(20\d{2})年\s*(\d{1,2})月\s*(\d{1,2})日", joined)
    if match:
        y, m, d = match.groups()
        return valid_date(f"{int(y):04d}-{int(m):02d}-{int(d):02d}")
    match = re.search(r"(20\d{6})", joined)
    if match:
        value = match.group(1)
        return valid_date(f"{value[:4]}-{value[4:6]}-{value[6:8]}")
    return ""


def extract_docx_text(path: Path) -> Tuple[str, str]:
    try:
        with ZipFile(path) as zf:
            xml = zf.read("word/document.xml")
        root = ET.fromstring(xml)
    except (BadZipFile, KeyError, ET.ParseError, OSError) as exc:
        return "", "error"

    paragraphs: List[str] = []
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    for paragraph in root.findall(".//w:p", namespace):
        pieces: List[str] = []
        for node in paragraph.iter():
            tag = node.tag.rsplit("}", 1)[-1]
            if tag == "t" and node.text:
                pieces.append(node.text)
            elif tag == "tab":
                pieces.append(" ")
        line = clean_text("".join(pieces))
        if line:
            paragraphs.append(line)
    return "\n".join(paragraphs), "ok"


def vision_ocr_binary(base_dir: Path = ROOT) -> Optional[Path]:
    if not shutil.which("clang"):
        return None
    source = base_dir / "agent" / "vision_ocr.m"
    if not source.exists():
        return None
    bin_dir = base_dir / WORK_DIR_NAME / "bin"
    bin_dir.mkdir(parents=True, exist_ok=True)
    binary = bin_dir / "vision_ocr"
    if not binary.exists() or source.stat().st_mtime > binary.stat().st_mtime:
        result = subprocess.run(
            [
                "clang",
                "-fobjc-arc",
                str(source),
                "-framework",
                "Foundation",
                "-framework",
                "Vision",
                "-framework",
                "ImageIO",
                "-framework",
                "CoreGraphics",
                "-o",
                str(binary),
            ],
            check=False,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=120,
        )
        if result.returncode != 0:
            return None
    return binary


def tesseract_languages() -> List[str]:
    if not shutil.which("tesseract"):
        return []
    result = subprocess.run(
        ["tesseract", "--list-langs"],
        check=False,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        encoding="utf-8",
        errors="ignore",
        timeout=15,
    )
    output = clean_text(result.stdout) + "\n" + clean_text(result.stderr)
    return [line.strip() for line in output.splitlines() if line.strip() and not line.startswith("List of")]


def tesseract_lang_arg() -> str:
    langs = set(tesseract_languages())
    choices = []
    for lang in ("chi_sim", "chi_tra", "eng"):
        if lang in langs:
            choices.append(lang)
    return "+".join(choices or ["eng"])


def render_pdf_pages(path: Path, tmp_dir: Path, max_pages: Optional[int] = None, dpi: int = DEFAULT_OCR_DPI) -> List[Path]:
    if not shutil.which("pdftoppm"):
        return []
    prefix = tmp_dir / "page"
    cmd = ["pdftoppm", "-r", str(dpi), "-png"]
    if max_pages:
        cmd.extend(["-f", "1", "-l", str(max_pages)])
    cmd.extend([str(path), str(prefix)])
    try:
        result = subprocess.run(
            cmd,
            check=False,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            encoding="utf-8",
            errors="ignore",
            timeout=300,
        )
    except subprocess.TimeoutExpired:
        return []
    if result.returncode != 0:
        return []
    return sorted(tmp_dir.glob("page-*.png"))


def ocr_images_with_tesseract(images: List[Path]) -> str:
    lang = tesseract_lang_arg()
    chunks: List[str] = []
    for image in images:
        result = subprocess.run(
            ["tesseract", str(image), "stdout", "-l", lang, "--psm", "6"],
            check=False,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            encoding="utf-8",
            errors="ignore",
            timeout=180,
        )
        if clean_text(result.stdout):
            chunks.append(result.stdout)
    return clean_text("\n\n".join(chunks))


def ocr_images_with_vision(images: List[Path], base_dir: Path = ROOT) -> str:
    binary = vision_ocr_binary(base_dir)
    if not binary:
        return ""
    result = subprocess.run(
        [str(binary), *[str(image) for image in images]],
        check=False,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        encoding="utf-8",
        errors="ignore",
        timeout=max(180, 60 * len(images)),
    )
    if result.returncode != 0 and not clean_text(result.stdout):
        return ""
    text = re.sub(r"^===PAGE:.*?===$", "", result.stdout, flags=re.MULTILINE)
    return clean_text(text)


def ocr_pdf_text(path: Path, max_pages: Optional[int] = None, base_dir: Path = ROOT) -> Tuple[str, str]:
    with tempfile.TemporaryDirectory(prefix="gongwen-ocr-") as tmp:
        images = render_pdf_pages(path, Path(tmp), max_pages=max_pages)
        if not images:
            return "", "needs_ocr"
        if shutil.which("tesseract"):
            text = ocr_images_with_tesseract(images)
            if len(text) >= MIN_USEFUL_TEXT:
                return text, "ok_ocr_tesseract"
        text = ocr_images_with_vision(images, base_dir=base_dir)
        if len(text) >= MIN_USEFUL_TEXT:
            return text, "ok_ocr_vision"
    return "", "needs_ocr"


def extract_pdf_text(path: Path, use_ocr: bool = False, ocr_max_pages: Optional[int] = None, base_dir: Path = ROOT) -> Tuple[str, str]:
    if not shutil.which("pdftotext"):
        return ocr_pdf_text(path, max_pages=ocr_max_pages, base_dir=base_dir) if use_ocr else ("", "needs_pdftotext")
    try:
        result = subprocess.run(
            ["pdftotext", str(path), "-"],
            check=False,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            encoding="utf-8",
            errors="ignore",
            timeout=60,
        )
    except Exception as exc:
        return "", "error"
    text = clean_text(result.stdout)
    if result.returncode != 0 and not text:
        return "", "error"
    if len(text) < MIN_USEFUL_TEXT:
        if use_ocr:
            ocr_text, ocr_status = ocr_pdf_text(path, max_pages=ocr_max_pages, base_dir=base_dir)
            if ocr_text:
                return ocr_text, ocr_status
        return text, "needs_ocr"
    return text, "ok"


def extract_text(path: Path, use_ocr: bool = False, ocr_max_pages: Optional[int] = None, base_dir: Path = ROOT) -> Tuple[str, str]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx_text(path)
    if suffix == ".pdf":
        return extract_pdf_text(path, use_ocr=use_ocr, ocr_max_pages=ocr_max_pages, base_dir=base_dir)
    return "", "unsupported"


def iter_source_files(base_dir: Path) -> Iterable[Path]:
    for path in sorted(base_dir.rglob("*")):
        if not path.is_file():
            continue
        if path.name.startswith("~$") or path.suffix.lower() not in SUPPORTED_EXTS:
            continue
        lower_parts = {p.lower() for p in path.parts}
        if lower_parts & IGNORED_PARTS:
            continue
        if WORK_DIR_NAME in path.parts:
            continue
        yield path


def cached_text_is_useful(path: Path, text_dir: Path) -> bool:
    text_file = text_dir / f"{stable_id(path)}.txt"
    if not text_file.exists():
        return False
    return text_file.stat().st_size >= MIN_USEFUL_TEXT


def source_priority(path: Path, text_dir: Path) -> Tuple[int, str]:
    if is_template_file(path):
        return (0, str(path))
    if cached_text_is_useful(path, text_dir):
        return (1, str(path))
    if path.suffix.lower() == ".docx":
        return (2, str(path))
    return (3, str(path))


def choose_coverage_files(source_files: List[Path], text_dir: Path, per_type: Optional[int]) -> List[Path]:
    if not per_type or per_type <= 0:
        return source_files
    groups: Dict[str, List[Path]] = {}
    for path in source_files:
        groups.setdefault(infer_business_from_path(path), []).append(path)
    chosen: List[Path] = []
    for business_type in sorted(groups):
        chosen.extend(sorted(groups[business_type], key=lambda item: source_priority(item, text_dir))[:per_type])
    return sorted(chosen)


def build_index(
    base_dir: Path = ROOT,
    force: bool = False,
    use_ocr: bool = False,
    ocr_max_pages: Optional[int] = None,
    coverage_per_type: Optional[int] = None,
) -> Dict[str, object]:
    work_dir = base_dir / WORK_DIR_NAME
    text_dir = work_dir / TEXT_DIR_NAME
    text_dir.mkdir(parents=True, exist_ok=True)
    records: List[FileRecord] = []

    all_source_files = list(iter_source_files(base_dir))
    source_files = choose_coverage_files(all_source_files, text_dir, coverage_per_type)
    total = len(source_files)
    for idx, path in enumerate(source_files, start=1):
        if use_ocr:
            print(f"[{idx}/{total}] 解析：{path.relative_to(base_dir)}", flush=True)
        file_id = stable_id(path)
        text_file = text_dir / f"{file_id}.txt"
        should_extract = True
        if text_file.exists() and not force:
            text = text_file.read_text(encoding="utf-8", errors="ignore")
            if use_ocr and path.suffix.lower() == ".pdf" and len(text) < MIN_USEFUL_TEXT:
                should_extract = True
            else:
                should_extract = False
                parse_status = "ok" if len(text) >= MIN_USEFUL_TEXT or path.suffix.lower() == ".docx" else "needs_ocr"
        if should_extract:
            text, parse_status = extract_text(path, use_ocr=use_ocr, ocr_max_pages=ocr_max_pages, base_dir=base_dir)
            text_file.write_text(text, encoding="utf-8")

        relative_path = path.relative_to(base_dir)
        title = title_from_name(path)
        region = infer_region(str(relative_path))
        record = FileRecord(
            id=file_id,
            path=str(relative_path),
            extension=path.suffix.lower().lstrip("."),
            business_type=infer_business_from_path(path),
            region=region,
            source_role=infer_source_role(path),
            is_template=is_template_file(path),
            title=title,
            project_name_guess=guess_project_name(title),
            issuing_org_guess=guess_issuing_org(title),
            date_guess=guess_date(text, path),
            text_chars=len(text),
            parse_status=parse_status,
            text_path=str(text_file.relative_to(base_dir)),
        )
        records.append(record)

    stats = Counter()
    for record in records:
        stats[f"ext:{record.extension}"] += 1
        stats[f"status:{record.parse_status}"] += 1
        stats[f"type:{record.business_type}"] += 1
        if record.is_template:
            stats["templates"] += 1
    if coverage_per_type:
        stats["coverage_per_type"] = coverage_per_type
        stats["source_files_total"] = len(all_source_files)

    index = {
        "version": 1,
        "base_dir": str(base_dir),
        "built_at": dt.datetime.now().isoformat(timespec="seconds"),
        "records": [asdict(record) for record in records],
        "stats": dict(sorted(stats.items())),
    }
    work_dir.mkdir(parents=True, exist_ok=True)
    (work_dir / INDEX_NAME).write_text(json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8")
    return index


def load_index(base_dir: Path = ROOT) -> Dict[str, object]:
    index_file = base_dir / WORK_DIR_NAME / INDEX_NAME
    if not index_file.exists():
        return build_index(base_dir)
    return json.loads(index_file.read_text(encoding="utf-8"))


def records_from_index(index: Dict[str, object]) -> List[FileRecord]:
    return [FileRecord(**item) for item in index.get("records", [])]


def read_record_text(record: FileRecord, base_dir: Path = ROOT, limit: int = 6000) -> str:
    path = base_dir / record.text_path
    if not path.exists():
        return ""
    return path.read_text(encoding="utf-8", errors="ignore")[:limit]


def char_ngrams(text: str, n: int = 2) -> Counter:
    text = clean_text(text)
    tokens: List[str] = []
    words = re.findall(r"[A-Za-z0-9_.+-]+|[\u4e00-\u9fff]", text)
    compact = "".join(words)
    for i in range(max(0, len(compact) - n + 1)):
        tokens.append(compact[i : i + n])
    for word in re.findall(r"[A-Za-z0-9_.+-]{2,}", text):
        tokens.append(word.lower())
    return Counter(tokens)


def cosine_like(a: Counter, b: Counter) -> float:
    if not a or not b:
        return 0.0
    common = set(a) & set(b)
    dot = sum(a[key] * b[key] for key in common)
    norm_a = sum(v * v for v in a.values()) ** 0.5
    norm_b = sum(v * v for v in b.values()) ** 0.5
    if not norm_a or not norm_b:
        return 0.0
    return dot / (norm_a * norm_b)


def classify_business_type(project: ProjectInput, records: List[FileRecord]) -> str:
    if project.business_type:
        for record in records:
            if project.business_type == record.business_type or project.business_type in record.business_type:
                return record.business_type
        return strip_business_prefix(project.business_type)

    text = " ".join([project.project_name, project.special_notes, project.approval_basis, project.scale])
    scores = Counter()
    for business_type, aliases in BUSINESS_ALIASES.items():
        for alias in aliases:
            if alias in text:
                scores[business_type] += 3
    if scores:
        return scores.most_common(1)[0][0]
    return "文物保护许可申请"


def retrieve(
    project: ProjectInput,
    records: List[FileRecord],
    business_type: str,
    top_k: int = DEFAULT_TOP_K,
    templates_only: bool = False,
) -> List[FileRecord]:
    query = " ".join(
        [
            business_type,
            project.region,
            project.project_name,
            project.construction_unit,
            project.location,
            project.scale,
            project.land_area,
            project.approval_basis,
            project.special_notes,
        ]
    )
    query_vec = char_ngrams(query)
    scored: List[Tuple[float, FileRecord]] = []
    for record in records:
        if templates_only and not record.is_template:
            continue
        if record.parse_status.startswith("error"):
            continue
        text = " ".join(
            [
                record.title,
                record.business_type,
                record.region,
                record.source_role,
                record.project_name_guess,
                read_record_text(record, limit=2000),
            ]
        )
        score = cosine_like(query_vec, char_ngrams(text))
        if record.business_type == business_type:
            score += 0.45
        elif business_type in record.business_type or record.business_type in business_type:
            score += 0.2
        if project.region and record.region == project.region:
            score += 0.25
        if record.is_template:
            score += 0.15 if templates_only else 0.04
        if record.parse_status == "needs_ocr":
            score -= 0.2
        if score > 0:
            scored.append((score, record))
    scored.sort(key=lambda item: item[0], reverse=True)
    return [record for _, record in scored[:top_k]]


def select_template(project: ProjectInput, records: List[FileRecord], business_type: str) -> Optional[FileRecord]:
    templates = retrieve(project, records, business_type, top_k=1, templates_only=True)
    if templates:
        return templates[0]
    fallback = [r for r in records if r.is_template and r.business_type == business_type]
    return fallback[0] if fallback else None


def extract_attachment_hints(template_text: str, business_type: str) -> List[str]:
    attachments: List[str] = []
    in_attachment = False
    for raw_line in template_text.splitlines():
        line = clean_text(raw_line)
        if not line:
            continue
        if line.startswith("附件"):
            in_attachment = True
            continue
        if in_attachment:
            match = re.match(r"^[（(]?\d+[）).、]?\s*(.+)", line)
            if match:
                item = clean_text(match.group(1))
                item = re.sub(r"\s+", "", item)
                if item and item not in attachments:
                    attachments.append(item)
            elif len(line) > 40:
                break
    return attachments or DEFAULT_ATTACHMENTS.get(business_type, DEFAULT_ATTACHMENTS["文物保护许可申请"])


def default_attachment_items(project: ProjectInput, business_type: str) -> List[str]:
    subject = project.project_name or "XX项目"
    items = DEFAULT_ATTACHMENTS.get(business_type, DEFAULT_ATTACHMENTS["文物保护许可申请"])
    normalized = [item.replace("XX项目", subject) for item in items]

    if business_type == "文物保护许可申请" and project.approval_basis:
        match = re.search(
            r"《([^》]+(?:复函|批复|核准文件|备案告知书|选址意见)[^》]*)》",
            project.approval_basis,
        )
        if match:
            normalized[0] = match.group(1)
    return normalized


def today_chinese() -> str:
    now = dt.date.today()
    return f"{now.year}年{now.month}月{now.day}日"


def make_title(project: ProjectInput, business_type: str) -> str:
    subject = project.project_name or "XX项目"
    org = project.issuing_org or project.construction_unit or "XX单位"
    if business_type == "文物保护许可申请":
        return f"{org}关于办理{subject}用地范围内文物保护许可的请示"
    if business_type == "勘探验收请示":
        return f"关于申请{subject}考古勘探验收的函"
    if business_type == "申请开展考古勘探工作请示":
        return f"{org}关于申请开展{subject}考古勘探工作的请示"
    if business_type == "发掘请示":
        return f"{org}关于{subject}用地范围内涉及文物遗迹进行考古发掘的请示"
    if business_type == "开工请示":
        return f"{org}关于{subject}选址范围内文物勘探工作验收后开工的请示"
    if business_type == "勘探计划备案请示":
        return f"{org}关于{subject}占地范围内文物勘探计划备案的请示"
    if business_type == "勘探报告备案请示":
        return f"{org}关于{subject}占地范围内文物勘探成果报告备案的请示"
    if business_type == "文物调查意见":
        return f"{org}关于{subject}选址范围内文物审查意见的函"
    if business_type == "文物调查报告":
        return f"关于{subject}文物调查的报告"
    if business_type == "文物保护安全责任书":
        return f"{subject}文物保护安全责任书"
    return f"{org}关于{subject}有关文物事宜的请示"


def build_body(project: ProjectInput, business_type: str, attachment_items: List[str]) -> List[str]:
    recipient = project.recipient_org or "XX文物主管部门"
    subject = project.project_name or "XX项目"
    unit = project.construction_unit or project.issuing_org or "建设单位"
    location = project.location or "项目建设地点"
    scale = project.scale
    area = project.land_area or "项目用地面积"
    basis = project.approval_basis or "相关审批、核准或备案文件"
    notes = project.special_notes

    if business_type == "勘探验收请示":
        unit_label = project.issuing_org or project.construction_unit or "我司"
        project_purpose = "为完善项目相关手续"
        third_party = "第三方公司"
        source_text = " ".join([project.scale, project.approval_basis, project.special_notes])
        report_name = f"《{subject}考古调查、勘探工作报告》"
        result_sentence = ""
        if not re.search(r"(未发现文化遗存|未发现.*遗迹|无遗迹|未见文物遗存)", source_text):
            match = re.search(r"(?:共)?发现([^。；;，,]*?(?:遗迹|遗存)[^。；;，,]*?(?:\\d+处|[一二三四五六七八九十百]+处))", source_text)
            if match:
                result_sentence = f"项目区域及周边范围内共发现{match.group(1)}。"
            elif re.search(r"(发现.*(?:遗迹|遗存))", source_text):
                result_sentence = "项目区域及周边范围内发现相关遗迹。"
        project_info = f"该项目位于{location}，项目面积为{area}"
        if scale:
            project_info += f"，建设内容为{scale}"
        project_info += "。"
        return [
            f"{recipient}：",
            f"{project_purpose}，{unit_label}已委托{third_party}开展{subject}考古调查、勘探工作。{project_info}",
            f"根据《内蒙古自治区文物局关于做好基本建设用地考古工作的通知》（内文物发〔2025〕6号）要求，{third_party}已完成项目所在区域实地考古调查、勘探工作，并编制形成{report_name}。{result_sentence}",
            "现将相关资料予以呈报，恳请贵局进行审查，并组织专家对勘探成果进行验收。",
            "此函。",
            "附件：",
            f"{subject}考古调查、勘探工作报告",
        ]

    if business_type == "申请开展考古勘探工作请示":
        unit_label = project.issuing_org or project.construction_unit or "我单位"
        plan_name = f"《{subject}考古调查、勘探工作计划》"
        intro = (
            f"{subject}由{unit_label}承担建设，项目位于{location}，项目面积为{area}。"
            f"为落实文物保护相关要求，查明该项目用地范围内地下文物遗存分布情况，"
            "做好项目实施前文物保护工作，现就开展考古勘探工作请示如下。"
        )
        request = (
            "依据《中华人民共和国文物保护法》及《内蒙古自治区文物保护条例》等有关规定，"
            f"{unit_label}已委托勘探单位编制完成{plan_name}。"
            "现申请对上述区域开展文物勘探工作，恳请贵局按程序报送材料、协调组织勘探工作。"
        )
        return [f"{recipient}：", intro, request, "妥否，请批示。"]

    if business_type == "文物调查报告":
        report_org = project.issuing_org or "调查单位"
        source_text = " ".join([project.approval_basis, project.special_notes]).strip()
        if project.approval_basis and project.approval_basis != "相关审批、核准或备案文件":
            opening = f"{project.approval_basis}已收悉。"
        else:
            opening = f"贵单位关于{subject}开展文物调查的来函已收悉。"
        project_info_parts = [f"{subject}位于{location}"]
        if scale:
            project_info_parts.append(f"建设内容为{scale}")
        if project.land_area:
            project_info_parts.append(f"项目用地面积为{area}")
        project_info = "，".join(project_info_parts) + "。"

        if notes:
            survey = f"经查阅相关文物资料并开展现场调查，{notes.rstrip('。')}。"
        else:
            survey = "经查阅相关文物资料并开展现场调查，项目用地范围内文物分布情况已完成核查。"

        if re.search(r"(发现|涉及).*(遗迹|遗存|文物|保护单位|疑点)", source_text) and not re.search(
            r"(未发现|无遗迹|未见)", source_text
        ):
            conclusion = "相关调查结论和保护要求应按现场调查材料及文物主管部门意见执行。"
        else:
            conclusion = (
                f"{report_org}原则同意此项目建设。根据相关规定，建议在此项目开工建设前，"
                "对该建设工程范围内有可能埋藏文物的地方进行考古勘探工作。"
                "如在施工建设工程中发现文物，应立即停工，并按程序报告当地文物行政管理部门。"
            )

        return [
            f"{recipient}：",
            f"{opening}{report_org}会同有关单位组成调查组，对{subject}进行了现场文物调查，具体情况如下：",
            project_info,
            survey,
            conclusion,
            "专此",
            "附件：",
            *attachment_items,
        ]

    if business_type == "发掘请示":
        unit_label = project.issuing_org or project.construction_unit or "我公司"
        project_info_parts = [f"{subject}位于{location}", f"项目面积为{area}"]
        if scale:
            project_info_parts.append(f"建设内容为{scale}")
        project_info = "，".join(project_info_parts) + "。"

        survey_sentence = "目前，我公司已委托第三方机构完成考古调查、勘探工作并形成相关成果材料。"
        if notes:
            clean_notes = re.sub(r"^(?:勘探后|经调查勘探[,，]?|经勘探[,，]?)", "", notes).strip(" ，,。")
            survey_sentence = (
                "目前，我公司已委托第三方机构完成考古调查、勘探工作并形成相关成果材料。"
                f"经调查勘探，{clean_notes}。"
            )

        basis_sentence = "依据《内蒙古自治区文物局关于做好基本建设用地考古工作的通知》（内文物发〔2025〕6号）相关规定"
        if "内文物发〔2025〕6号" not in basis and basis != "相关审批、核准或备案文件":
            basis_sentence = f"依据{basis}及《内蒙古自治区文物局关于做好基本建设用地考古工作的通知》（内文物发〔2025〕6号）相关规定"

        return [
            f"{recipient}：",
            f"{unit_label}正在办理{subject}相关手续，{project_info}",
            survey_sentence,
            (
                f"{basis_sentence}，为保障项目合规推进及文物保护工作落实，"
                "现申请对上述遗迹开展考古发掘工作。恳请贵局向相关部门报送材料，协调组织考古发掘工作。"
            ),
            "特此请示。",
            "附件：",
            f"{subject}考古调查、勘探工作报告",
        ]

    if business_type == "文物保护许可申请":
        intro = (
            f"{subject}由{unit}承担建设，项目位于{location}。"
            "为依法依规办理项目用地范围内文物保护相关手续，现就该项目用地范围内文物保护许可事宜请示如下。"
        )
    else:
        intro = (
            f"{subject}由{unit}承担建设，项目位于{location}。"
            f"目前该项目已取得{basis}。为依法依规做好基本建设工程涉及文物保护相关工作，现就有关事项请示如下。"
        )
    paragraphs = [f"{recipient}：", intro]
    number_names = "一二三四五六七八九十"
    section_items = [
        ("项目名称", f"{subject}。"),
    ]
    if scale:
        section_items.append(("建设规模及主要建设内容", f"{scale}。"))
    section_items.extend(
        [
            ("建设地点", f"{location}。"),
            ("项目用地面积", f"{subject}拟占地面积为{area}。"),
        ]
    )
    for idx, (heading, content) in enumerate(section_items):
        paragraphs.extend([f"{number_names[idx]}、{heading}", content])

    if business_type == "文物保护许可申请":
        request = (
            "按照文物保护相关法律法规及基本建设用地考古工作要求，"
            "现申请贵局对该项目用地范围内是否涉及文物保护单位、不可移动文物及其保护范围、"
            "建设控制地带等情况进行核查，并请出具相关审查意见。"
        )
    elif business_type == "勘探验收请示":
        request = (
            f"{subject}用地范围内文物勘探工作已按要求组织实施，相关成果材料已形成。"
            "现申请对文物勘探工作成果进行验收，并请予以审查批复。"
        )
    elif business_type == "发掘请示":
        request = (
            f"根据前期调查勘探情况，{subject}用地范围内涉及需开展考古发掘的文物遗迹。"
            "为妥善保护地下文物，现申请组织开展考古发掘工作，请予审查。"
        )
    elif business_type == "开工请示":
        request = (
            f"{subject}选址范围内文物勘探工作已完成并通过相关验收。"
            "为保障项目建设进度，现申请在落实文物保护要求的前提下开展后续建设工作。"
        )
    elif business_type == "勘探计划备案请示":
        request = (
            f"为推进{subject}文物勘探工作，现将项目文物勘探计划及相关资料报送备案，请予审查。"
        )
    elif business_type == "勘探报告备案请示":
        request = (
            f"{subject}文物勘探成果报告已编制完成，现将成果报告及相关资料报送备案，请予审查。"
        )
    elif business_type == "文物调查意见":
        request = (
            f"经对{subject}选址范围开展文物调查核查，相关意见拟按程序出具。请结合项目实际落实文物保护要求。"
        )
    elif business_type == "文物保护安全责任书":
        request = (
            f"{unit}承诺在{subject}实施过程中严格落实文物保护主体责任，"
            "发现文物遗存或疑似文物线索时立即停止相关作业并按程序报告。"
        )
    else:
        request = f"现将{subject}有关文物事宜报请审查，请予支持。"

    paragraphs.append(request)
    if notes:
        paragraphs.extend(["五、其他需要说明的事项", notes])
    paragraphs.extend(["妥否，请批示。", "附件："])
    for idx, item in enumerate(attachment_items, start=1):
        paragraphs.append(f"{idx}. {item}")
    return paragraphs


def build_markdown(
    project: ProjectInput,
    business_type: str,
    template: Optional[FileRecord],
    references: List[FileRecord],
    body: List[str],
    checks: List[Dict[str, str]],
    warnings: List[str],
) -> str:
    title = make_title(project, business_type)
    footer_lines = [
        "",
        project.issuing_org or project.construction_unit or "XX单位",
        today_chinese(),
    ]
    contact_parts = []
    if project.contact:
        contact_parts.append(f"联系人：{project.contact}")
    if project.phone:
        contact_parts.append(f"电话：{project.phone}")
    if contact_parts:
        footer_lines.extend(["", "    ".join(contact_parts)])

    source_lines = ["## 参考来源"]
    if template:
        source_lines.append(f"- 模板：{template.path}")
    for record in references:
        source_lines.append(f"- 参考件：{record.path}")

    check_lines = ["## 事实校验"]
    for item in checks:
        mark = "通过" if item["status"] == "pass" else "待补充"
        check_lines.append(f"- {mark}：{item['field']} - {item['message']}")
    for warning in warnings:
        check_lines.append(f"- 提醒：{warning}")

    draft = ["# " + title, ""]
    draft.extend(body)
    draft.extend(footer_lines)
    draft.extend(["", "---", ""])
    draft.extend(source_lines)
    draft.extend([""])
    draft.extend(check_lines)
    return "\n\n".join(draft) + "\n"


def make_checks(project: ProjectInput, draft_text: str) -> List[Dict[str, str]]:
    required = [
        ("项目名称", project.project_name),
        ("建设单位", project.construction_unit),
        ("建设地点", project.location),
        ("用地面积", project.land_area),
        ("发文主体", project.issuing_org),
        ("主送机关", project.recipient_org),
    ]
    checks: List[Dict[str, str]] = []
    for field_name, value in required:
        if not value:
            checks.append({"field": field_name, "status": "warn", "message": "用户未填写，已保留通用表述或占位。"})
        elif value in draft_text:
            checks.append({"field": field_name, "status": "pass", "message": "已在草稿中原样出现。"})
        else:
            checks.append({"field": field_name, "status": "warn", "message": f"未在草稿中原样检出：{value}"})
    return checks


def set_chinese_font(run, font_name: str = "仿宋_GB2312", size: Optional[int] = None) -> None:
    if Pt is None:
        return
    run.font.name = font_name
    if size:
        run.font.size = Pt(size)
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_docx_paragraph(
    doc,
    text: str,
    *,
    align=None,
    bold: bool = False,
    size: int = 16,
    first_line: bool = True,
    hanging_indent: Optional[int] = None,
):
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = Pt(28)
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    if hanging_indent is not None:
        paragraph_format.left_indent = Pt(hanging_indent)
        paragraph_format.first_line_indent = Pt(-hanging_indent)
    elif first_line:
        paragraph_format.first_line_indent = Pt(size * 2)
    run = paragraph.add_run(text)
    run.bold = bold
    set_chinese_font(run, size=size)
    return paragraph


def add_docx_attachment_item(doc, index: int, text: str):
    paragraph = doc.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = Pt(28)
    paragraph_format.left_indent = Pt(16 * 7)
    if index == 1:
        paragraph_format.first_line_indent = Pt(16 * 2) - Pt(16 * 7)
        line = f"附件：1.\t{text}"
    else:
        paragraph_format.first_line_indent = Pt(16 * 5) - Pt(16 * 7)
        line = f"{index}.\t{text}"
    paragraph_format.tab_stops.clear_all()
    paragraph_format.tab_stops.add_tab_stop(Pt(16 * 7))
    run = paragraph.add_run(line)
    set_chinese_font(run, size=16)
    return paragraph


def estimate_text_width_pt(text: str, size: int = 16) -> float:
    width = 0.0
    for char in text or "":
        width += size * 0.55 if char.isascii() else size
    return width


def signature_center_position(date_text: str):
    content_width = Cm(21 - 2.8 - 2.6)
    date_right_edge_pt = content_width.pt - Pt(16 * 4).pt
    return Pt(max(0, date_right_edge_pt - estimate_text_width_pt(date_text) / 2))


def add_docx_signature(doc, issuing_org: str, date_text: str):
    doc.add_paragraph()

    signature = doc.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.LEFT
    signature.paragraph_format.left_indent = Pt(0)
    signature.paragraph_format.right_indent = Pt(0)
    signature.paragraph_format.space_before = Pt(0)
    signature.paragraph_format.space_after = Pt(0)
    signature.paragraph_format.line_spacing = Pt(28)
    signature.paragraph_format.tab_stops.clear_all()
    signature.paragraph_format.tab_stops.add_tab_stop(
        signature_center_position(date_text),
        WD_TAB_ALIGNMENT.CENTER,
    )
    run = signature.add_run("\t" + issuing_org)
    set_chinese_font(run, size=16)

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date.paragraph_format.right_indent = Pt(16 * 4)
    date.paragraph_format.space_before = Pt(0)
    date.paragraph_format.space_after = Pt(0)
    date.paragraph_format.line_spacing = Pt(28)
    run = date.add_run(date_text)
    set_chinese_font(run, size=16)


def create_docx(project: ProjectInput, business_type: str, body: List[str], checks: List[Dict[str, str]], output_path: Path) -> None:
    if Document is None:
        raise RuntimeError("缺少 python-docx，无法生成 Word 文件。")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    normal = doc.styles["Normal"]
    normal.font.name = "仿宋_GB2312"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    normal.font.size = Pt(16)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = Pt(28)

    title = make_title(project, business_type)
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    set_chinese_font(title_run, font_name="方正小标宋简体", size=22)
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(24)
    title_para.paragraph_format.line_spacing = Pt(30)

    in_attachments = False
    attachment_index = 1
    for line in body:
        if line == "附件：":
            in_attachments = True
            continue
        if in_attachments:
            item = re.sub(r"^\d+[.、]\s*", "", line).strip()
            if item:
                add_docx_attachment_item(doc, attachment_index, item)
                attachment_index += 1
        elif re.match(r"^[一二三四五六七八九十]+、", line):
            add_docx_paragraph(doc, line, bold=True, first_line=False)
        elif line.endswith("：") and len(line) < 30:
            add_docx_paragraph(doc, line, first_line=False)
        else:
            add_docx_paragraph(doc, line)

    add_docx_signature(doc, project.issuing_org or project.construction_unit or "XX单位", today_chinese())
    contact_parts = []
    if project.contact:
        contact_parts.append(f"联系人：{project.contact}")
    if project.phone:
        contact_parts.append(f"电话：{project.phone}")
    if contact_parts:
        add_docx_paragraph(doc, "    ".join(contact_parts), first_line=False)

    doc.save(output_path)


def generate_draft(project: ProjectInput, base_dir: Path = ROOT) -> DraftResult:
    index = load_index(base_dir)
    records = records_from_index(index)
    business_type = classify_business_type(project, records)
    template = select_template(project, records, business_type)
    candidates = retrieve(project, records, business_type, top_k=DEFAULT_TOP_K * 3, templates_only=False)
    candidates = [
        record
        for record in candidates
        if record.business_type == business_type or business_type in record.business_type or record.business_type in business_type
    ]
    usable_refs = [record for record in candidates if not record.is_template and record.parse_status == "ok"]
    title_only_refs = [record for record in candidates if not record.is_template and record.parse_status == "needs_ocr"]
    references = (usable_refs + title_only_refs)[:DEFAULT_TOP_K]

    template_text = read_record_text(template, base_dir=base_dir) if template else ""
    attachments = extract_attachment_hints(template_text, business_type)
    body = build_body(project, business_type, attachments)
    draft_preview = "\n".join([make_title(project, business_type), *body])
    checks = make_checks(project, draft_preview)

    warnings: List[str] = []
    if not template:
        warnings.append("未找到同类模板，已按内置公文结构生成初稿。")
    if any(record.parse_status == "needs_ocr" for record in records):
        warnings.append("部分 PDF 为扫描件或不可提取文字，需安装 OCR 后重建索引以提升参考质量。")

    output_dir = default_output_dir(base_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    timestamp = dt.datetime.now().strftime("%Y%m%d-%H%M%S")
    basename = safe_filename(f"{timestamp} {project.project_name or '公文初稿'}")
    docx_path = output_dir / f"{basename}.docx"

    create_docx(project, business_type, body, checks, docx_path)

    return DraftResult(
        classified_business_type=business_type,
        selected_template=template,
        references=references,
        docx_path=result_path_for_display(docx_path, base_dir),
        checks=checks,
        warnings=warnings,
    )


def print_index_summary(index: Dict[str, object]) -> None:
    print(f"索引时间：{index.get('built_at')}")
    print(f"资料库：{index.get('base_dir')}")
    print(f"文件数：{len(index.get('records', []))}")
    print("统计：")
    for key, value in index.get("stats", {}).items():
        print(f"  {key}: {value}")


def parse_project_json(path: Path) -> ProjectInput:
    data = json.loads(path.read_text(encoding="utf-8"))
    return ProjectInput.from_mapping(data)


def run_cli(argv: Optional[List[str]] = None) -> int:
    parser = argparse.ArgumentParser(description="本地专属公文撰写 Agent")
    sub = parser.add_subparsers(dest="command", required=True)

    p_index = sub.add_parser("index", help="扫描资料库并建立索引")
    p_index.add_argument("--force", action="store_true", help="强制重新解析所有文件")
    p_index.add_argument("--ocr", action="store_true", help="对扫描 PDF 启用 OCR；优先 tesseract，缺失时使用 macOS Vision")
    p_index.add_argument("--ocr-max-pages", type=int, default=None, help="每个 PDF 最多 OCR 前 N 页；默认识别全部页面")
    p_index.add_argument("--coverage-per-type", type=int, default=None, help="每个公文类型最多入库 N 个文件；优先模板和已解析文件")

    p_search = sub.add_parser("search", help="检索模板和历史参考件")
    p_search.add_argument("query", help="检索关键词")
    p_search.add_argument("--type", default="", help="事项类型")
    p_search.add_argument("--region", default="", help="旗区")
    p_search.add_argument("-k", type=int, default=DEFAULT_TOP_K)

    p_draft = sub.add_parser("draft", help="根据项目 JSON 生成桌面格式化 Word")
    p_draft.add_argument("project_json", type=Path)

    p_web = sub.add_parser("web", help="启动本地 Web 表单")
    p_web.add_argument("--host", default="127.0.0.1")
    p_web.add_argument("--port", type=int, default=8765)

    args = parser.parse_args(argv)
    if args.command == "index":
        index = build_index(
            ROOT,
            force=args.force,
            use_ocr=args.ocr,
            ocr_max_pages=args.ocr_max_pages,
            coverage_per_type=args.coverage_per_type,
        )
        print_index_summary(index)
        return 0
    if args.command == "search":
        index = load_index(ROOT)
        records = records_from_index(index)
        project = ProjectInput(project_name=args.query, business_type=args.type, region=args.region)
        business_type = classify_business_type(project, records)
        for idx, record in enumerate(retrieve(project, records, business_type, top_k=args.k), 1):
            print(f"{idx}. [{record.business_type}] {record.title}")
            print(f"   {record.path}")
            print(f"   旗区={record.region or '-'} 模板={record.is_template} 状态={record.parse_status}")
        return 0
    if args.command == "draft":
        project = parse_project_json(args.project_json)
        result = generate_draft(project, ROOT)
        print(json.dumps(asdict(result), ensure_ascii=False, indent=2))
        return 0
    if args.command == "web":
        ensure_index()
        server = ThreadingHTTPServer((args.host, args.port), AgentHandler)
        print(f"本地公文撰写 Agent 已启动：http://{args.host}:{args.port}")
        try:
            server.serve_forever()
        except KeyboardInterrupt:
            print("\n已停止")
        return 0
    return 1


def ensure_index() -> None:
    if not (ROOT / WORK_DIR_NAME / INDEX_NAME).exists():
        build_index(ROOT)


FORM_FIELDS = [
    ("business_type", "事项类型", "如：文物保护许可申请、勘探验收请示；可留空自动判断"),
    ("issuing_org", "发文主体", "如：XX公司、XX文物局"),
    ("recipient_org", "主送机关", "如：准格尔旗文物局、鄂尔多斯市文物局"),
    ("project_name", "项目名称", "必须准确填写"),
    ("construction_unit", "建设单位", "必须准确填写"),
    ("location", "建设地点", "如：鄂尔多斯市准格尔旗大路镇..."),
    ("scale", "建设规模及主要内容", "线路长度、装机容量、建设内容等"),
    ("land_area", "用地面积", "如：0.1148公顷"),
    ("approval_basis", "审批/核准依据", "项目备案、核准、批复文件名称"),
    ("contact", "联系人", "可选"),
    ("phone", "电话", "可选"),
    ("special_notes", "特殊说明", "需额外写入正文的情况"),
]


def render_form(message: str = "", result: Optional[DraftResult] = None) -> bytes:
    ensure_index()
    index = load_index(ROOT)
    stats = index.get("stats", {})
    field_html = []
    for name, label, placeholder in FORM_FIELDS:
        if name in {"scale", "special_notes"}:
            field_html.append(
                f'<label>{label}<textarea name="{name}" placeholder="{html.escape(placeholder)}"></textarea></label>'
            )
        else:
            field_html.append(
                f'<label>{label}<input name="{name}" placeholder="{html.escape(placeholder)}"></label>'
            )
    result_html = ""
    if result:
        checks = "".join(
            f"<li>{html.escape(item['field'])}：{html.escape(item['message'])}</li>" for item in result.checks
        )
        refs = "".join(f"<li>{html.escape(record.path)}</li>" for record in result.references)
        result_html = f"""
        <section class="result">
          <h2>已生成格式化 Word</h2>
          <p>识别事项类型：<strong>{html.escape(result.classified_business_type)}</strong></p>
          <p><a href="/file?path={urllib.parse.quote(result.docx_path)}">下载格式化 Word</a></p>
          <h3>事实校验</h3><ul>{checks}</ul>
          <h3>参考件</h3><ul>{refs or "<li>未检索到参考件</li>"}</ul>
        </section>
        """
    page = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <title>专属公文撰写 Agent</title>
  <style>
    body {{ margin: 0; font-family: -apple-system, BlinkMacSystemFont, "PingFang SC", "Microsoft YaHei", sans-serif; background:#f7f7f5; color:#1f2933; }}
    main {{ max-width: 1040px; margin: 0 auto; padding: 32px 24px 56px; }}
    header {{ margin-bottom: 24px; }}
    h1 {{ font-size: 28px; margin: 0 0 8px; }}
    .meta {{ color:#5b6472; font-size: 14px; }}
    form {{ display:grid; grid-template-columns: repeat(2, minmax(0, 1fr)); gap:16px; background:white; border:1px solid #ddded8; padding:20px; border-radius:8px; }}
    label {{ display:flex; flex-direction:column; gap:7px; font-size:14px; font-weight:600; }}
    input, textarea {{ border:1px solid #cfd4dc; border-radius:6px; padding:10px 12px; font: inherit; font-weight:400; background:#fff; }}
    textarea {{ min-height: 92px; resize: vertical; }}
    label:nth-last-child(1), label:nth-last-child(2) {{ grid-column: span 2; }}
    button {{ justify-self:start; border:0; border-radius:6px; background:#245b45; color:white; padding:10px 18px; font:inherit; font-weight:700; cursor:pointer; }}
    .message, .result {{ margin: 18px 0; padding: 16px 18px; border-radius:8px; background:#fff; border:1px solid #ddded8; }}
    .result a {{ display:inline-block; margin-right:12px; color:#175c3d; font-weight:700; }}
    ul {{ line-height:1.8; }}
    @media (max-width: 760px) {{ form {{ grid-template-columns: 1fr; }} label:nth-last-child(1), label:nth-last-child(2) {{ grid-column: span 1; }} }}
  </style>
</head>
<body>
<main>
  <header>
    <h1>专属公文撰写 Agent</h1>
    <div class="meta">已入库 {len(index.get("records", []))} 个文件；模板 {stats.get("templates", 0)} 个；需 OCR 的 PDF {stats.get("status:needs_ocr", 0)} 个。</div>
  </header>
  {f'<div class="message">{html.escape(message)}</div>' if message else ''}
  {result_html}
  <form method="post" action="/draft">
    {''.join(field_html)}
    <button type="submit">生成格式化 Word</button>
  </form>
</main>
</body>
</html>"""
    return page.encode("utf-8")


class AgentHandler(BaseHTTPRequestHandler):
    def do_GET(self) -> None:  # noqa: N802
        parsed = urllib.parse.urlparse(self.path)
        if parsed.path == "/":
            self.send_html(render_form())
            return
        if parsed.path == "/file":
            params = urllib.parse.parse_qs(parsed.query)
            rel = params.get("path", [""])[0]
            self.serve_file(rel)
            return
        self.send_error(HTTPStatus.NOT_FOUND)

    def do_POST(self) -> None:  # noqa: N802
        parsed = urllib.parse.urlparse(self.path)
        if parsed.path != "/draft":
            self.send_error(HTTPStatus.NOT_FOUND)
            return
        length = int(self.headers.get("Content-Length", "0"))
        body = self.rfile.read(length).decode("utf-8", errors="ignore")
        values = {key: vals[0] for key, vals in urllib.parse.parse_qs(body).items()}
        try:
            result = generate_draft(ProjectInput.from_mapping(values), ROOT)
            self.send_html(render_form(result=result))
        except Exception as exc:
            self.send_html(render_form(message=f"生成失败：{exc}"), status=500)

    def send_html(self, content: bytes, status: int = 200) -> None:
        self.send_response(status)
        self.send_header("Content-Type", "text/html; charset=utf-8")
        self.send_header("Content-Length", str(len(content)))
        self.end_headers()
        self.wfile.write(content)

    def serve_file(self, rel_path: str) -> None:
        path = (ROOT / rel_path).resolve()
        output_root = (ROOT / WORK_DIR_NAME / OUTPUT_DIR_NAME).resolve()
        desktop_root = (Path.home() / "Desktop").resolve()
        allowed_roots = [output_root, desktop_root]
        if not any(str(path).startswith(str(root)) for root in allowed_roots) or not path.exists():
            self.send_error(HTTPStatus.NOT_FOUND)
            return
        content = path.read_bytes()
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if path.suffix.lower() == ".md":
            content_type = "text/markdown; charset=utf-8"
        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Disposition", f"attachment; filename*=UTF-8''{urllib.parse.quote(path.name)}")
        self.send_header("Content-Length", str(len(content)))
        self.end_headers()
        self.wfile.write(content)

    def log_message(self, fmt: str, *args) -> None:
        sys.stderr.write("%s - %s\n" % (self.address_string(), fmt % args))


if __name__ == "__main__":
    raise SystemExit(run_cli())
